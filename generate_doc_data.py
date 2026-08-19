# -*- coding: utf-8 -*-
"""
Full definition of 72 Use Cases for DevCine Appendix A.
Generates docs/Dac_Ta_Use_Case.docx
"""
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from build_full_use_cases import add_use_case_table, format_run

use_cases = [
    # ==========================================
    # KHỐI 1: TÁC NHÂN KHÁCH HÀNG (UC-01 -> UC-22)
    # ==========================================
    {
        "num": 1,
        "id": "UC-01",
        "name": "Đăng ký tài khoản khách hàng mới",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Cho phép người dùng mới tạo tài khoản thành viên trong hệ thống để tham gia đặt vé trực tuyến, tích lũy điểm thưởng và nhận các ưu đãi thành viên.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang đăng ký tài khoản từ giao diện trang chủ hoặc màn hình đăng nhập.",
            "Bước 2: Khách hàng nhập các thông tin bắt buộc gồm họ tên, số điện thoại, địa chỉ email và mật khẩu bảo mật.",
            "Bước 3: Khách hàng nhấn nút \"Đăng ký\".",
            "Bước 4: Hệ thống kiểm tra tính hợp lệ của dữ liệu (định dạng email, định dạng số điện thoại, độ mạnh mật khẩu và đảm bảo số điện thoại/email chưa từng được đăng ký trước đó).",
            "Bước 5: Hệ thống khởi tạo tài khoản thành viên mới với hạng mặc định là Đồng (Bronze), tự động đăng nhập và chuyển hướng khách hàng về trang chủ."
        ],
        "notes": [
            "Nếu số điện thoại hoặc email đã tồn tại trong hệ thống, hệ thống hiển thị thông báo lỗi và yêu cầu kiểm tra lại.",
            "Mật khẩu phải đáp ứng tiêu chuẩn an toàn từ 8 đến 32 ký tự, bao gồm chữ hoa, chữ thường, số và ký tự đặc biệt."
        ]
    },
    {
        "num": 2,
        "id": "UC-02",
        "name": "Đăng nhập hệ thống",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Xác thực danh tính khách hàng để cấp quyền truy cập các tính năng cá nhân hóa như đặt vé, xem lịch sử mua hàng và quản lý ưu đãi.",
        "steps": [
            "Bước 1: Khách hàng chọn chức năng \"Đăng nhập\" trên thanh điều hướng.",
            "Bước 2: Khách hàng nhập định danh tài khoản (Số điện thoại hoặc Email) và mật khẩu.",
            "Bước 3: Khách hàng nhấn nút \"Đăng nhập\".",
            "Bước 4: Hệ thống kiểm tra thông tin tài khoản và xác thực mật khẩu.",
            "Bước 5: Hệ thống lưu trạng thái phiên đăng nhập, tải thông tin thành viên (điểm thưởng, hạng thẻ) và chuyển về trang đang thao tác trước đó."
        ],
        "notes": [
            "Nếu thông tin định danh hoặc mật khẩu không chính xác, hệ thống thông báo lỗi và không tiết lộ cụ thể trường thông tin nào sai để đảm bảo an toàn bảo mật."
        ]
    },
    {
        "num": 3,
        "id": "UC-03",
        "name": "Khôi phục mật khẩu qua Email",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng tự khôi phục quyền truy cập vào tài khoản khi bị quên mật khẩu thông qua mã xác minh (OTP) gửi về hộp thư điện tử.",
        "steps": [
            "Bước 1: Tại màn hình đăng nhập, khách hàng chọn \"Quên mật khẩu?\".",
            "Bước 2: Khách hàng nhập địa chỉ email đã đăng ký tài khoản và nhấn \"Gửi mã xác nhận\".",
            "Bước 3: Hệ thống kiểm tra email trong hệ thống, tạo mã xác thực ngẫu nhiên gồm 6 chữ số có hiệu lực giới hạn và gửi email đến khách hàng.",
            "Bước 4: Khách hàng kiểm tra hộp thư, nhập mã xác thực vào giao diện.",
            "Bước 5: Hệ thống xác thực mã thành công và cho phép khách hàng nhập mật khẩu mới cùng xác nhận mật khẩu.",
            "Bước 6: Khách hàng nhấn \"Đặt lại mật khẩu\", hệ thống lưu mật khẩu mới và thông báo thành công."
        ],
        "notes": [
            "Mã xác thực chỉ có hiệu lực trong thời gian ngắn và chỉ được sử dụng một lần duy nhất.",
            "Giao diện có cơ chế đếm ngược thời gian chờ gửi lại mã để chống tình trạng gửi thư liên tục."
        ]
    },
    {
        "num": 4,
        "id": "UC-04",
        "name": "Xem và cập nhật thông tin cá nhân",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng xem hồ sơ cá nhân, cấp bậc thành viên, tổng điểm tích lũy và cập nhật thông tin liên lạc khi có nhu cầu thay đổi.",
        "steps": [
            "Bước 1: Khách hàng đăng nhập và chọn mục \"Hồ sơ cá nhân\" từ menu tài khoản.",
            "Bước 2: Hệ thống hiển thị thông tin chi tiết: Họ tên, Email, Số điện thoại, Hạng thành viên hiện tại và Điểm thưởng tích lũy.",
            "Bước 3: Khách hàng chỉnh sửa các trường thông tin cần cập nhật (Họ tên, Số điện thoại, Email).",
            "Bước 4: Khách hàng nhấn nút \"Lưu thay đổi\".",
            "Bước 5: Hệ thống kiểm tra tính hợp lệ và duy nhất của thông tin mới, lưu vào cơ sở dữ liệu và hiển thị thông báo cập nhật thành công."
        ],
        "notes": [
            "Không cho phép sửa đổi số điện thoại hoặc email trùng lặp với tài khoản của khách hàng khác đang hoạt động."
        ]
    },
    {
        "num": 5,
        "id": "UC-05",
        "name": "Đổi mật khẩu tài khoản",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng chủ động thay đổi mật khẩu đăng nhập định kỳ để nâng cao tính an toàn cho tài khoản cá nhân.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang \"Hồ sơ cá nhân\" và chọn tab \"Đổi mật khẩu\".",
            "Bước 2: Khách hàng nhập Mật khẩu hiện tại, Mật khẩu mới và Nhập lại mật khẩu mới.",
            "Bước 3: Khách hàng nhấn nút \"Cập nhật mật khẩu\".",
            "Bước 4: Hệ thống kiểm tra tính chính xác của mật khẩu hiện tại, độ mạnh của mật khẩu mới và sự trùng khớp của ô xác nhận mật khẩu.",
            "Bước 5: Hệ thống mã hóa và lưu mật khẩu mới, sau đó thông báo đổi mật khẩu thành công."
        ],
        "notes": [
            "Nếu mật khẩu cũ không đúng hoặc mật khẩu mới trùng khớp hoàn toàn với mật khẩu cũ, hệ thống sẽ đưa ra thông báo cảnh báo."
        ]
    },
    {
        "num": 6,
        "id": "UC-06",
        "name": "Xem danh sách phim đang chiếu và sắp chiếu",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Cung cấp danh sách các bộ phim đang được công chiếu tại các rạp và những bộ phim chuẩn bị ra mắt giúp khách hàng dễ dàng nắm bắt thông tin giải trí.",
        "steps": [
            "Bước 1: Khách hàng truy cập vào Trang chủ hoặc mục \"Lịch chiếu\" trên thanh điều hướng.",
            "Bước 2: Khách hàng chuyển đổi giữa các tab \"Phim đang chiếu\" và \"Phim sắp chiếu\".",
            "Bước 3: Hệ thống hiển thị danh sách phim kèm áp-phích (poster), tên phim, thể loại, thời lượng, phân loại độ tuổi và điểm đánh giá trung bình."
        ],
        "notes": [
            "Chỉ các bộ phim đang ở trạng thái kích hoạt công khai mới được hiển thị trên giao diện người dùng."
        ]
    },
    {
        "num": 7,
        "id": "UC-07",
        "name": "Xem thông tin chi tiết phim và xem Video Trailer",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Hiển thị toàn bộ thông tin chi tiết của một bộ phim (nội dung tóm tắt, đạo diễn, diễn viên, ngày khởi chiếu) và cho phép xem video giới thiệu (trailer).",
        "steps": [
            "Bước 1: Khách hàng nhấn chọn vào một bộ phim từ danh sách phim hoặc trang chủ.",
            "Bước 2: Hệ thống hiển thị trang chi tiết phim với đầy đủ thông tin: Áp-phích, Tên phim, Tên gốc, Thể loại, Thời lượng, Giới hạn độ tuổi, Đạo diễn, Diễn viên, Tóm tắt nội dung và các suất chiếu theo từng rạp.",
            "Bước 3: Khách hàng nhấn vào nút \"Xem Trailer\".",
            "Bước 4: Hệ thống mở cửa sổ phát video trailer chính thức của phim để khách hàng thưởng thức."
        ],
        "notes": [
            "Hệ thống cảnh báo rõ ràng nhãn kiểm duyệt độ tuổi (như P, K, T13, T16, T18, C) để khách hàng cân nhắc trước khi quyết định đặt vé."
        ]
    },
    {
        "num": 8,
        "id": "UC-08",
        "name": "Tìm kiếm và lọc phim",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Hỗ trợ khách hàng tìm kiếm nhanh bộ phim mong muốn theo từ khóa tên phim hoặc lọc phim theo thể loại, cụm rạp và định dạng chiếu.",
        "steps": [
            "Bước 1: Khách hàng nhập từ khóa tên phim vào thanh tìm kiếm hoặc truy cập trang \"Tìm kiếm\".",
            "Bước 2: Khách hàng có thể chọn thêm các tiêu chí lọc: Thể loại phim (Hành động, Hài, Hoạt hình...), Cụm rạp chiếu, Định dạng (2D, 3D, IMAX).",
            "Bước 3: Hệ thống tự động xử lý trì hoãn tìm kiếm (debounce) và trả về danh sách các bộ phim phù hợp với tiêu chí đã chọn."
        ],
        "notes": [
            "Nếu không tìm thấy kết quả phù hợp, hệ thống hiển thị trạng thái danh sách rỗng kèm gợi ý các phim đang được quan tâm nhiều nhất."
        ]
    },
    {
        "num": 9,
        "id": "UC-09",
        "name": "Tra cứu Lịch chiếu và chọn Suất chiếu theo Cụm rạp",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng xem lịch chiếu tổng hợp của tất cả các rạp, chọn ngày xem, chọn cụm rạp yêu thích và chọn suất chiếu thuận tiện nhất.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang \"Lịch chiếu\" hoặc xem phần lịch chiếu tại trang Chi tiết phim.",
            "Bước 2: Khách hàng chọn Cụm rạp (theo khu vực/tỉnh thành) và chọn Ngày chiếu mong muốn.",
            "Bước 3: Hệ thống lọc và hiển thị danh sách các phim kèm các khung giờ chiếu, phòng chiếu và định dạng tương ứng.",
            "Bước 4: Khách hàng chọn một suất chiếu cụ thể để bắt đầu quá trình đặt vé."
        ],
        "notes": [
            "Các suất chiếu đã diễn ra trong quá khứ sẽ tự động được ẩn đi hoặc khóa không cho chọn đặt vé."
        ]
    },
    {
        "num": 10,
        "id": "UC-10",
        "name": "Chọn vị trí ghế ngồi trên Sơ đồ phòng chiếu",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Hiển thị sơ đồ mặt bằng ghế của phòng chiếu theo thời gian thực, cho phép khách hàng chọn các vị trí ghế ngồi ưng ý và tạm giữ chỗ trong 10 phút.",
        "steps": [
            "Bước 1: Sau khi chọn suất chiếu, hệ thống điều hướng khách hàng đến màn hình Sơ đồ ghế.",
            "Bước 2: Hệ thống hiển thị trực quan cấu trúc phòng chiếu gồm Màn hình, Lối đi và Ma trận các loại ghế (Ghế thường, Ghế VIP, Ghế đôi Sweetbox) kèm trạng thái (Còn trống, Đang giữ, Đã bán, Đang bảo trì).",
            "Bước 3: Khách hàng nhấn chọn các vị trí ghế mong muốn và chọn loại đối tượng áp dụng cho từng ghế (Người lớn, Học sinh/Sinh viên).",
            "Bước 4: Khách hàng nhấn \"Tiếp tục\". Hệ thống thực hiện khóa giữ chỗ tạm thời các ghế đã chọn trong thời gian 10 phút để khách tiến hành thanh toán."
        ],
        "notes": [
            "Hệ thống kiểm tra quy tắc không để lại ghế trống đơn lẻ (orphan seat) ở giữa hoặc đầu hàng.",
            "Nếu có ghế đã bị người khác chọn trước trong tích tắc, hệ thống sẽ thông báo xung đột và yêu cầu khách hàng chọn vị trí khác."
        ]
    },
    {
        "num": 11,
        "id": "UC-11",
        "name": "Lựa chọn Bắp nước / Combo và tùy biến lựa chọn",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng chọn thêm các món bắp rang, nước ngọt, thức ăn nhẹ hoặc combo ưu đãi kèm theo vé, đồng thời hỗ trợ đổi vị và chọn kích cỡ.",
        "steps": [
            "Bước 1: Tại bước chọn bắp nước của quy trình đặt vé, hệ thống hiển thị danh mục các combo và món ăn đang mở bán kèm hình ảnh, mô tả và giá tiền.",
            "Bước 2: Khách hàng tăng/giảm số lượng món hoặc combo muốn mua.",
            "Bước 3: Đối với các combo có tùy chọn vị (như Vị bắp: Phô mai, Caramel; Vị nước: Coca, Sprite), hệ thống mở cửa sổ tùy biến để khách hàng lựa chọn khẩu vị yêu thích và tính phụ thu tương ứng (nếu có).",
            "Bước 4: Hệ thống tự động cập nhật tổng tiền giỏ hàng bao gồm tiền vé và tiền bắp nước."
        ],
        "notes": [
            "Bước chọn bắp nước là không bắt buộc; khách hàng có thể bấm \"Bỏ qua\" hoặc \"Tiếp tục\" để sang bước thanh toán."
        ]
    },
    {
        "num": 12,
        "id": "UC-12",
        "name": "Áp dụng Mã giảm giá / Voucher ưu đãi vào đơn đặt vé",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Hỗ trợ khách hàng nhập mã khuyến mãi hoặc chọn các voucher sẵn có trong ví ưu đãi để được khấu trừ giảm giá trực tiếp trên tổng hóa đơn.",
        "steps": [
            "Bước 1: Tại bước xác nhận thanh toán, khách hàng nhấn chọn danh sách voucher khả dụng từ ví cá nhân hoặc nhập mã code khuyến mãi vào ô nhập liệu.",
            "Bước 2: Khách hàng nhấn nút \"Áp dụng\".",
            "Bước 3: Hệ thống kiểm tra điều kiện áp dụng của mã: Thời hạn sử dụng, giá trị đơn hàng tối thiểu, giới hạn số vé, phim áp dụng và đối tượng khách hàng.",
            "Bước 4: Nếu thỏa mãn điều kiện, hệ thống tính toán số tiền được giảm trừ và cập nhật lại số tiền thanh toán thực tế của đơn hàng."
        ],
        "notes": [
            "Mỗi đơn hàng chỉ được áp dụng một mã giảm giá trừ khi chương trình có quy định cho phép cộng dồn ưu đãi.",
            "Nếu mã giảm giá không đủ điều kiện áp dụng, hệ thống hiển thị lý do cụ thể (chưa đủ tiền tối thiểu, hết hạn...)."
        ]
    },
    {
        "num": 13,
        "id": "UC-13",
        "name": "Thanh toán vé trực tuyến qua Cổng thanh toán VNPAY",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Kết nối cổng thanh toán trực tuyến an toàn VNPAY cho phép khách hàng thanh toán qua ứng dụng ngân hàng, ví điện tử hoặc thẻ nội địa/quốc tế.",
        "steps": [
            "Bước 1: Khách hàng kiểm tra lại tóm tắt đơn đặt vé (Phim, Suất chiếu, Ghế ngồi, Combo, Tổng tiền) và chọn phương thức thanh toán VNPAY.",
            "Bước 2: Khách hàng nhấn nút \"Thanh toán\".",
            "Bước 3: Hệ thống tạo giao dịch và chuyển hướng khách hàng sang cổng thanh toán bảo mật VNPAY.",
            "Bước 4: Khách hàng thực hiện quét mã QR qua ứng dụng ngân hàng hoặc nhập thông tin thẻ để xác nhận thanh toán.",
            "Bước 5: Sau khi thanh toán thành công, VNPAY chuyển hướng phản hồi về hệ thống DevCine.",
            "Bước 6: Hệ thống xác thực chữ ký bảo mật giao dịch, chuyển trạng thái đơn hàng sang Đã xác nhận (Confirmed), hoàn tất đặt chỗ và cộng điểm thưởng thành viên."
        ],
        "notes": [
            "Nếu khách hàng hủy thanh toán hoặc giao dịch thất bại quá thời gian giữ chỗ 10 phút, hệ thống sẽ tự động giải phóng các ghế đã chọn về trạng thái trống."
        ]
    },
    {
        "num": 14,
        "id": "UC-14",
        "name": "Xem Vé điện tử (Mã đặt vé & Mã QR Code) sau khi thanh toán",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Cung cấp vé điện tử trực quan chứa thông tin suất chiếu và mã vạch QR Code giúp khách hàng sử dụng để check-in hoặc in vé giấy khi đến rạp.",
        "steps": [
            "Bước 1: Khi thanh toán thành công, hệ thống điều hướng khách hàng đến màn hình \"Đặt vé thành công\".",
            "Bước 2: Hệ thống hiển thị toàn bộ thông tin vé điện tử: Mã đặt vé, Tên phim, Rạp chiếu, Phòng chiếu, Thời gian chiếu, Vị trí ghế, Danh sách bắp nước và Mã QR Code xác thực.",
            "Bước 3: Hệ thống đồng thời gửi email thông báo xác nhận kèm hóa đơn chi tiết đến địa chỉ email của khách hàng."
        ],
        "notes": [
            "Khách hàng có thể chụp màn hình hoặc lưu mã QR để xuất trình trực tiếp cho nhân viên soát vé tại rạp chiếu phim."
        ]
    },
    {
        "num": 15,
        "id": "UC-15",
        "name": "Tra cứu Lịch sử đặt vé và trạng thái vé",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng xem lại toàn bộ danh sách các đơn hàng và vé xem phim đã mua trong quá khứ hoặc sắp tới kèm thông tin chi tiết từng vé.",
        "steps": [
            "Bước 1: Khách hàng đăng nhập và chọn mục \"Lịch sử đặt vé\" trong hồ sơ cá nhân.",
            "Bước 2: Hệ thống truy vấn và hiển thị danh sách các đơn đặt vé theo thứ tự thời gian từ mới nhất đến cũ nhất.",
            "Bước 3: Khách hàng nhấn chọn vào một đơn đặt vé bất kỳ để xem chi tiết: Mã đặt vé, Tên phim, Suất chiếu, Ghế ngồi, Bắp nước, Tổng tiền thanh toán, Trạng thái (Đã thanh toán, Đã in vé/Check-in, Đã hủy) và Mã QR vé."
        ],
        "notes": [
            "Dữ liệu lịch sử được tối ưu hóa để tải nhanh chóng và bảo toàn đầy đủ thông tin kể cả khi giá vé hoặc thực đơn có sự thay đổi sau này."
        ]
    },
    {
        "num": 16,
        "id": "UC-16",
        "name": "Đánh giá và chấm điểm phim",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cho phép những khách hàng đã thực sự mua vé và xem phim gửi đánh giá chấm điểm sao cùng cảm nhận bình luận về bộ phim để chia sẻ với cộng đồng.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang chi tiết của bộ phim đã từng xem.",
            "Bước 2: Hệ thống kiểm tra điều kiện (khách hàng đã đăng nhập và có đơn đặt vé thành công bộ phim này).",
            "Bước 3: Khách hàng chọn số sao đánh giá (từ 1 đến 5 sao) và nhập nội dung nhận xét cảm nghĩ.",
            "Bước 4: Khách hàng nhấn nút \"Gửi đánh giá\".",
            "Bước 5: Hệ thống ghi nhận đánh giá, tính toán lại điểm đánh giá trung bình của phim và hiển thị nhận xét trên trang phim."
        ],
        "notes": [
            "Khách hàng chưa mua vé bộ phim này sẽ không được cấp quyền gửi đánh giá nhằm ngăn ngừa tình trạng đánh giá ảo hoặc tiêu cực vô căn cứ."
        ]
    },
    {
        "num": 17,
        "id": "UC-17",
        "name": "Quản lý Ví Voucher cá nhân",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cung cấp giao diện ví voucher cá nhân để khách hàng tra cứu mã khuyến mãi, lưu mã ưu đãi vào tài khoản và theo dõi hạn sử dụng của các voucher đang sở hữu.",
        "steps": [
            "Bước 1: Khách hàng truy cập mục \"Ưu đãi của tôi\" trong trang cá nhân.",
            "Bước 2: Hệ thống hiển thị các tab phân loại: \"Voucher khả dụng\", \"Đổi điểm lấy ưu đãi\" và \"Lịch sử voucher (Đã dùng / Hết hạn)\".",
            "Bước 3: Khách hàng có thể nhập mã code bí mật được nhận từ các chiến dịch quảng cáo vào ô tra cứu và nhấn \"Lưu mã\".",
            "Bước 4: Hệ thống kiểm tra tính hợp lệ của mã và thêm voucher vào ví cá nhân của khách hàng."
        ],
        "notes": [
            "Nếu mã ưu đãi đã hết lượt sử dụng, hết hạn hoặc khách hàng đã lưu trước đó, hệ thống sẽ đưa ra thông báo phù hợp."
        ]
    },
    {
        "num": 18,
        "id": "UC-18",
        "name": "Quy đổi Điểm tích lũy thành viên lấy Voucher giảm giá",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cho phép khách hàng sử dụng điểm thưởng tích lũy (Loyalty Points) có được từ các lần mua vé trước đó để đổi lấy các phiếu giảm giá có giá trị cao.",
        "steps": [
            "Bước 1: Khách hàng truy cập tab \"Đổi điểm lấy ưu đãi\" tại trang Ví Voucher.",
            "Bước 2: Hệ thống hiển thị danh sách các chương trình khuyến mãi cho phép đổi điểm kèm số điểm yêu cầu tương ứng (ví dụ: 100 điểm, 200 điểm) và số điểm hiện có của khách.",
            "Bước 3: Khách hàng nhấn nút \"Đổi ngay\" tại gói ưu đãi mong muốn.",
            "Bước 4: Hệ thống kiểm tra số điểm tích lũy hiện tại của khách hàng có đủ điều kiện không.",
            "Bước 5: Hệ thống trừ số điểm tương ứng trong tài khoản thành viên, sinh voucher mới vào ví của khách hàng và thông báo đổi điểm thành công."
        ],
        "notes": [
            "Mỗi gói ưu đãi quy đổi chỉ có thể đổi một lần cho mỗi tài khoản thành viên theo chính sách của từng chương trình."
        ]
    },
    {
        "num": 19,
        "id": "UC-19",
        "name": "Gửi Yêu cầu hỗ trợ / Liên hệ trực tuyến",
        "priority": "Thấp",
        "actor": "Khách hàng",
        "description": "Cung cấp biểu mẫu liên hệ trực tuyến để khách hàng gửi phản ánh, thắc mắc về dịch vụ, sự cố vé hoặc góp ý đến ban quản trị cụm rạp.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang \"Liên hệ / Hỗ trợ\".",
            "Bước 2: Khách hàng chọn nhóm chủ đề liên hệ (Vấn đề về vé, Thẻ thành viên, Góp ý dịch vụ, Hợp tác quảng cáo).",
            "Bước 3: Khách hàng nhập số điện thoại và nội dung chi tiết cần hỗ trợ.",
            "Bước 4: Khách hàng nhấn nút \"Gửi yêu cầu hỗ trợ\".",
            "Bước 5: Hệ thống tiếp nhận, tạo phiếu yêu cầu hỗ trợ (Support Ticket) ở trạng thái chờ xử lý và thông báo gửi yêu cầu thành công đến khách hàng."
        ],
        "notes": [
            "Khách hàng cần đăng nhập tài khoản để hệ thống tự động liên kết yêu cầu hỗ trợ với hồ sơ khách hàng nhằm tiện cho việc phản hồi."
        ]
    },
    {
        "num": 20,
        "id": "UC-20",
        "name": "Xem và tra cứu Danh mục Câu hỏi thường gặp (FAQ)",
        "priority": "Thấp",
        "actor": "Khách hàng",
        "description": "Cung cấp danh sách các câu hỏi thường gặp và câu trả lời chuẩn xác được sắp xếp theo từng chủ đề giúp khách hàng tự tra cứu thông tin nhanh chóng.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang \"Hỏi đáp (FAQ)\".",
            "Bước 2: Hệ thống hiển thị các câu hỏi được phân nhóm theo danh mục (Đặt vé online, Quy định độ tuổi, Thẻ thành viên, Giá vé & Ưu đãi).",
            "Bước 3: Khách hàng có thể nhập từ khóa vào ô tìm kiếm nhanh.",
            "Bước 4: Khách hàng nhấn vào tiêu đề câu hỏi để mở rộng nội dung câu trả lời chi tiết."
        ],
        "notes": [
            "Giao diện hỗ trợ mở/đóng mượt mà giúp người dùng dễ dàng theo dõi nhiều nội dung mà không bị rối mắt."
        ]
    },
    {
        "num": 21,
        "id": "UC-21",
        "name": "Xem Tin tức, Bài viết sự kiện và Chương trình khuyến mãi",
        "priority": "Thấp",
        "actor": "Khách hàng",
        "description": "Giúp khách hàng theo dõi các bài viết tin tức điện ảnh, sự kiện ra mắt phim và các bài viết giới thiệu chương trình ưu đãi hấp dẫn đang diễn ra.",
        "steps": [
            "Bước 1: Khách hàng truy cập trang \"Khuyến mãi\" hoặc xem mục tin tức trên trang chủ.",
            "Bước 2: Hệ thống hiển thị danh sách các bài viết tin tức sự kiện kèm hình ảnh đại diện, tiêu đề và tóm tắt ngắn.",
            "Bước 3: Khách hàng nhấn chọn một bài viết cụ thể.",
            "Bước 4: Hệ thống hiển thị trang nội dung bài viết chi tiết cùng thời gian áp dụng của sự kiện."
        ],
        "notes": [
            "Chỉ các bài viết đang ở trạng thái kích hoạt và còn trong thời gian hiệu lực mới được hiển thị công khai."
        ]
    },
    {
        "num": 22,
        "id": "UC-22",
        "name": "Nhận và quản lý Thông báo cá nhân trên hệ thống",
        "priority": "Thấp",
        "actor": "Khách hàng",
        "description": "Hộp thư thông báo nội bộ giúp khách hàng nhận các cập nhật quan trọng như: xác nhận đặt vé thành công, nhận voucher quà tặng, thăng hạng thành viên.",
        "steps": [
            "Bước 1: Khách hàng đăng nhập và nhấn vào biểu tượng Chuông thông báo trên thanh điều hướng.",
            "Bước 2: Hệ thống hiển thị số lượng thông báo chưa đọc cùng danh sách các thông báo mới nhất.",
            "Bước 3: Khách hàng nhấn vào một thông báo để xem nội dung chi tiết; hệ thống tự động chuyển trạng thái thông báo đó sang \"Đã đọc\".",
            "Bước 4: Khách hàng có thể nhấn nút \"Đánh dấu tất cả đã đọc\" để xóa huy hiệu số tin chưa đọc."
        ],
        "notes": [
            "Hệ thống đảm bảo tính bảo mật, người dùng chỉ có thể xem và thao tác trên thông báo thuộc về chính tài khoản của mình."
        ]
    },

    # =========================================================================
    # KHỐI 2: TÁC NHÂN NHÂN VIÊN QUẦY & SOÁT VÉ (POS / CHECK-IN) (UC-23 -> UC-41)
    # =========================================================================
    {
        "num": 23,
        "id": "UC-23",
        "name": "Đăng nhập hệ thống Bán vé & Vận hành tại quầy (POS)",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Xác thực danh tính và quyền hạn của nhân viên để truy cập vào màn hình bán vé tại quầy (POS) và màn hình soát vé theo cụm rạp được phân công.",
        "steps": [
            "Bước 1: Nhân viên truy cập cổng đăng nhập nội bộ dành cho nhân sự.",
            "Bước 2: Nhân viên nhập tên đăng nhập và mật khẩu được cấp.",
            "Bước 3: Nhân viên nhấn nút \"Đăng nhập\".",
            "Bước 4: Hệ thống xác thực thông tin, kiểm tra vai trò và quyền hạn thao tác tại quầy, đồng thời tải phạm vi cụm rạp (Cinema Scoping) mà nhân viên trực thuộc.",
            "Bước 5: Hệ thống điều hướng nhân viên vào màn hình Bán vé POS hoặc màn hình Soát vé tương ứng."
        ],
        "notes": [
            "Nhân viên chỉ có quyền thao tác trên dữ liệu và suất chiếu thuộc cụm rạp mà mình được phân bổ công tác; mọi hành vi thao tác chéo rạp đều bị chặn."
        ]
    },
    {
        "num": 24,
        "id": "UC-24",
        "name": "Đổi mật khẩu bắt buộc trong lần đầu tiên đăng nhập",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cơ chế bảo mật bắt buộc nhân viên mới phải thay đổi mật khẩu mặc định do quản trị viên cấp ngay trong lần đăng nhập đầu tiên trước khi sử dụng hệ thống.",
        "steps": [
            "Bước 1: Nhân viên đăng nhập thành công với tài khoản được quản trị viên khởi tạo lần đầu.",
            "Bước 2: Hệ thống nhận diện cờ yêu cầu đổi mật khẩu lần đầu và tự động chuyển hướng khóa vào màn hình \"Đổi mật khẩu lần đầu\".",
            "Bước 3: Nhân viên nhập Mật khẩu mới và Nhập lại mật khẩu mới đáp ứng tiêu chuẩn an toàn.",
            "Bước 4: Nhân viên nhấn nút \"Xác nhận đổi mật khẩu\".",
            "Bước 5: Hệ thống cập nhật mật khẩu mới, gỡ bỏ cờ bắt buộc và mở khóa toàn bộ quyền hạn để nhân viên bắt đầu làm việc."
        ],
        "notes": [
            "Nhân viên không thể truy cập bất kỳ tính năng bán vé hay quản trị nào khác nếu chưa hoàn tất bước đổi mật khẩu lần đầu này."
        ]
    },
    {
        "num": 25,
        "id": "UC-25",
        "name": "Tra cứu lịch chiếu và tình trạng suất chiếu nhanh tại quầy",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cung cấp giao diện trực quan hiển thị danh sách các phim, phòng chiếu và suất chiếu trong ngày tại rạp giúp nhân viên tư vấn nhanh cho khách hàng.",
        "steps": [
            "Bước 1: Nhân viên mở màn hình Bán vé POS.",
            "Bước 2: Hệ thống tự động tải và hiển thị danh sách tất cả các suất chiếu từ ngày hôm nay trở đi của cụm rạp hiện tại.",
            "Bước 3: Nhân viên chọn nhanh ngày chiếu qua các tab (Hôm nay, Ngày mai...) hoặc chọn phim cụ thể.",
            "Bước 4: Hệ thống gom nhóm hiển thị theo từng bộ phim, kèm định dạng chiếu (2D, 3D), phòng chiếu và thời gian bắt đầu chiếu."
        ],
        "notes": [
            "Các suất chiếu đã bắt đầu quá thời gian cho phép bán vé trễ (mặc định 30 phút sau giờ chiếu) sẽ tự động được ẩn khỏi màn hình bán vé."
        ]
    },
    {
        "num": 26,
        "id": "UC-26",
        "name": "Bán vé xem phim cho khách vãng lai",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Hỗ trợ nhân viên chọn suất chiếu, chọn vị trí ghế ngồi trên sơ đồ, chọn loại vé phù hợp cho khách mua trực tiếp tại quầy không cần tài khoản.",
        "steps": [
            "Bước 1: Nhân viên chọn một suất chiếu theo yêu cầu của khách hàng.",
            "Bước 2: Hệ thống hiển thị sơ đồ ghế của phòng chiếu theo thời gian thực.",
            "Bước 3: Nhân viên nhấn chọn các vị trí ghế ngồi mà khách hàng yêu cầu.",
            "Bước 4: Nhân viên chỉ định loại đối tượng cho từng ghế (Người lớn, Học sinh/Sinh viên, Trẻ em, Người cao tuổi) để hệ thống áp dụng mức giá tương ứng.",
            "Bước 5: Nhân viên xác nhận thông tin và chuyển sang bước tiếp theo."
        ],
        "notes": [
            "Nhân viên có quyền đặc biệt cho phép bỏ qua quy tắc kiểm tra ghế trống đơn lẻ nếu khách hàng có yêu cầu đặc thù và được quản lý chấp thuận."
        ]
    },
    {
        "num": 27,
        "id": "UC-27",
        "name": "Bán kèm hoặc Bán nhanh Bắp nước độc lập tại quầy",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Hỗ trợ nhân viên bán bắp rang, nước uống kèm theo đơn mua vé hoặc bán nhanh các sản phẩm bắp nước cho khách vãng lai chỉ có nhu cầu mua ẩm thực.",
        "steps": [
            "Bước 1: Tại màn hình POS, nhân viên có thể chọn thêm combo bắp nước vào đơn vé hoặc chọn chế độ \"Bán nhanh Bắp nước (Concession Only)\".",
            "Bước 2: Hệ thống hiển thị danh mục thực đơn F&B phân loại theo Combo, Bắp rang, Nước uống, Đồ ăn vặt.",
            "Bước 3: Nhân viên chọn món, số lượng và tùy chọn vị bắp/loại nước theo yêu cầu của khách hàng.",
            "Bước 4: Hệ thống tự động tính toán tổng tiền hàng và đưa vào giỏ thanh toán."
        ],
        "notes": [
            "Đơn bán bắp nước độc lập không gắn liền với suất chiếu hay ghế ngồi, cho phép thanh toán và xuất hóa đơn ngay tức thì."
        ]
    },
    {
        "num": 28,
        "id": "UC-28",
        "name": "Tra cứu và định danh Khách hàng thành viên qua Số điện thoại để tích điểm",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cho phép nhân viên tra cứu nhanh thông tin thành viên bằng số điện thoại để áp dụng ưu đãi theo hạng thẻ và tích lũy điểm thưởng khi mua hàng tại quầy.",
        "steps": [
            "Bước 1: Nhân viên hỏi số điện thoại của khách hàng và nhập vào ô \"Tra cứu thành viên\" tại màn hình POS.",
            "Bước 2: Nhân viên nhấn nút \"Tra cứu\" hoặc phím Enter.",
            "Bước 3: Hệ thống tìm kiếm và hiển thị thông tin thành viên: Họ tên khách hàng, Hạng thẻ (Đồng, Bạc, Vàng, Bạch Kim) và Điểm thưởng hiện có.",
            "Bước 4: Hệ thống liên kết tài khoản thành viên vào đơn hàng hiện tại để tích điểm sau khi giao dịch hoàn tất."
        ],
        "notes": [
            "Nếu số điện thoại chưa được đăng ký trong hệ thống, nhân viên có thể tiến hành bán vé dưới dạng khách vãng lai thông thường."
        ]
    },
    {
        "num": 29,
        "id": "UC-29",
        "name": "Áp dụng Mã giảm giá / Voucher tại quầy cho khách thành viên",
        "priority": "Trung bình",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Hỗ trợ áp dụng các phiếu giảm giá, voucher quà tặng hoặc mã ưu đãi từ chương trình khuyến mãi cho khách hàng khi thanh toán tại quầy vé.",
        "steps": [
            "Bước 1: Sau khi đã định danh khách hàng thành viên, nhân viên mở danh sách voucher mà khách đang sở hữu hoặc nhập mã ưu đãi trực tiếp.",
            "Bước 2: Nhân viên chọn voucher phù hợp và nhấn \"Áp dụng\".",
            "Bước 3: Hệ thống kiểm tra điều kiện áp dụng và tự động khấu trừ số tiền giảm giá vào tổng giá trị đơn hàng."
        ],
        "notes": [
            "Chỉ các voucher hợp lệ, chưa sử dụng và còn trong thời hạn hiệu lực mới được chấp nhận áp dụng."
        ]
    },
    {
        "num": 30,
        "id": "UC-30",
        "name": "Thanh toán bằng Tiền mặt và tính tiền thừa trả khách",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Xử lý giao dịch thanh toán bằng tiền mặt tại quầy, tự động tính số tiền thừa cần trả lại khách và hỗ trợ quy tắc làm tròn tiền lẻ.",
        "steps": [
            "Bước 1: Nhân viên chọn phương thức thanh toán là \"Tiền mặt (CASH)\".",
            "Bước 2: Hệ thống hiển thị tổng số tiền cần thu. Đối với khách vãng lai, hệ thống tự động làm tròn số tiền về bội số 1.000 đồng gần nhất.",
            "Bước 3: Nhân viên nhập số tiền mặt khách hàng đưa.",
            "Bước 4: Hệ thống tự động tính toán và hiển thị số tiền thừa (tiền thối) cần trả lại cho khách.",
            "Bước 5: Nhân viên thu tiền, trả lại tiền thừa cho khách và nhấn nút \"Hoàn tất thanh toán\"."
        ],
        "notes": [
            "Số tiền làm tròn được lưu vết chuẩn xác để đảm bảo quá trình đối soát doanh thu cuối ca làm việc khớp hoàn toàn với số tiền thực tế trong két."
        ]
    },
    {
        "num": 31,
        "id": "UC-31",
        "name": "Thanh toán Chuyển khoản ngân hàng qua Mã QR tự động (VietQR)",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Sinh mã phản hồi nhanh VietQR động chứa chính xác số tài khoản, số tiền và nội dung đơn hàng giúp khách quét mã chuyển khoản nhanh tại quầy.",
        "steps": [
            "Bước 1: Nhân viên chọn phương thức thanh toán là \"Chuyển khoản (TRANSFER)\".",
            "Bước 2: Hệ thống tự động tạo mã VietQR động hiển thị trên màn hình phụ hướng về phía khách hàng.",
            "Bước 3: Khách hàng sử dụng ứng dụng ngân hàng trên điện thoại quét mã QR để thực hiện chuyển khoản.",
            "Bước 4: Sau khi nhân viên xác nhận tài khoản rạp đã nhận được tiền, nhân viên nhấn nút \"Xác nhận đã nhận tiền\".",
            "Bước 5: Hệ thống chuyển trạng thái đơn hàng sang Đã xác nhận và tiến hành xuất vé."
        ],
        "notes": [
            "Mã QR động chứa sẵn nội dung chuyển khoản là mã đơn hàng để hỗ trợ đối soát giao dịch ngân hàng chính xác."
        ]
    },
    {
        "num": 32,
        "id": "UC-32",
        "name": "In vé giấy và Biên lai hóa đơn thanh toán tại quầy",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Tự động xuất lệnh in vé xem phim dạng giấy nhiệt và biên lai hóa đơn bán hàng cho khách hàng ngay sau khi thanh toán thành công tại quầy.",
        "steps": [
            "Bước 1: Khi đơn hàng được thanh toán thành công, hệ thống mở cửa sổ mẫu in vé chuẩn.",
            "Bước 2: Mẫu in bao gồm đầy đủ thông tin: Tên cụm rạp, Tên phim, Phòng chiếu, Định dạng, Ngày giờ chiếu, Vị trí ghế, Chi tiết bắp nước, Tổng tiền thanh toán và Mã QR xác thực của đơn vé.",
            "Bước 3: Máy in nhiệt tại quầy tự động in vé giấy và biên lai để nhân viên trao cho khách hàng."
        ],
        "notes": [
            "Hệ thống cho phép cấu hình in tự động hoặc nhấn nút in thủ công theo thiết lập của từng quầy bán vé."
        ]
    },
    {
        "num": 33,
        "id": "UC-33",
        "name": "Lưu tạm đơn hàng và phục hồi giỏ hàng chờ tại quầy",
        "priority": "Trung bình",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cho phép nhân viên lưu tạm một đơn hàng đang chọn dở dang khi khách hàng cần thêm thời gian lựa chọn, để chuyển sang phục vụ khách tiếp theo mà không làm mất trạng thái ghế.",
        "steps": [
            "Bước 1: Khi khách hàng cần chờ bạn hoặc chưa quyết định xong, nhân viên nhấn nút \"Lưu tạm đơn hàng\".",
            "Bước 2: Hệ thống tạo mã đơn hàng tạm, giữ ghế trên hệ thống và chuyển giao diện POS về trạng thái sẵn sàng đón khách mới.",
            "Bước 3: Khi khách hàng quay lại, nhân viên mở danh sách \"Đơn hàng chờ\".",
            "Bước 4: Nhân viên chọn đơn hàng tương ứng và nhấn \"Mở lại đơn\".",
            "Bước 5: Hệ thống tải lại toàn bộ ghế ngồi và bắp nước đã chọn trước đó để tiếp tục thực hiện thanh toán."
        ],
        "notes": [
            "Đơn hàng lưu tạm có thời gian hết hạn tự động; nếu quá hạn mà chưa thanh toán, hệ thống sẽ tự động hủy đơn và giải phóng ghế."
        ]
    },
    {
        "num": 34,
        "id": "UC-34",
        "name": "Soát vé tự động bằng Camera quét mã QR",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Sử dụng camera hoặc đầu đọc mã vạch để quét mã QR trên vé điện tử của khách hàng nhằm kiểm tra tính hợp lệ và check-in vào phòng chiếu.",
        "steps": [
            "Bước 1: Nhân viên soát vé mở màn hình \"Soát vé (Check-in)\" và kích hoạt Camera quét mã.",
            "Bước 2: Khách hàng xuất trình mã QR trên điện thoại hoặc vé giấy trước ống kính camera.",
            "Bước 3: Hệ thống nhận diện mã QR, phát âm thanh báo hiệu và hiển thị thông tin vé: Tên phim, Phòng chiếu, Suất chiếu, Vị trí ghế ngồi và Cảnh báo kiểm tra thẻ sinh viên/độ tuổi (nếu có).",
            "Bước 4: Nhân viên kiểm tra khách hàng thực tế và nhấn \"Xác nhận vào phòng / In vé\".",
            "Bước 5: Hệ thống cập nhật trạng thái vé thành \"Đã soát vé / Đã in vé\" và ghi nhận thời gian cùng nhân viên thực hiện."
        ],
        "notes": [
            "Nếu mã QR không hợp lệ, đã bị hủy hoặc vé đã được quét check-in trước đó, hệ thống sẽ lập tức phát âm thanh cảnh báo và hiển thị thông báo lỗi chống quét trùng lặp."
        ]
    },
    {
        "num": 35,
        "id": "UC-35",
        "name": "Soát vé và xác minh đơn thủ công bằng Mã đặt vé",
        "priority": "Trung bình",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Hỗ trợ nhân viên nhập mã ký tự đặt vé bằng tay trong trường hợp điện thoại của khách hàng bị hỏng camera, mờ mã QR hoặc lỗi thiết bị quét.",
        "steps": [
            "Bước 1: Nhân viên chuyển sang tab \"Nhập mã thủ công\" tại màn hình Soát vé.",
            "Bước 2: Nhân viên hỏi mã đặt vé (ví dụ: BK123456) của khách và nhập vào ô tìm kiếm.",
            "Bước 3: Nhân viên nhấn nút \"Tra cứu & Kiểm tra\".",
            "Bước 4: Hệ thống tìm kiếm và hiển thị chi tiết thông tin đơn vé.",
            "Bước 5: Nhân viên đối chiếu và xác nhận check-in đơn vé cho khách vào phòng chiếu."
        ],
        "notes": [
            "Chỉ các đơn vé thuộc cụm rạp mà nhân viên đang phụ trách mới có thể thực hiện kiểm tra và check-in thành công."
        ]
    },
    {
        "num": 36,
        "id": "UC-36",
        "name": "Tra cứu hóa đơn và in vé giấy từ mã đặt vé trực tuyến",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cho phép khách hàng đã mua vé online đến quầy xuất trình mã đặt vé để nhân viên in ra thành vé giấy truyền thống trước khi vào rạp xem phim.",
        "steps": [
            "Bước 1: Nhân viên nhập mã đặt vé online của khách hàng vào hệ thống tại quầy soát vé/in vé.",
            "Bước 2: Hệ thống kiểm tra đơn hàng: Đã thanh toán (Confirmed) và Chưa từng in vé giấy.",
            "Bước 3: Nhân viên nhấn nút \"In vé giấy\".",
            "Bước 4: Máy in xuất vé giấy, hệ thống cập nhật thời gian in vé và ghi nhận người thực hiện in.",
            "Bước 5: Hệ thống đồng thời gửi email thông báo cảm ơn đến khách hàng."
        ],
        "notes": [
            "Hệ thống kiểm soát nghiêm ngặt: Mỗi đơn đặt vé chỉ được in vé giấy một lần duy nhất để chống gian lận in lậu vé."
        ]
    },
    {
        "num": 37,
        "id": "UC-37",
        "name": "Tra cứu thông tin vé và người ngồi ghế khi xảy ra sự cố phòng chiếu",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Hỗ trợ nhân viên và quản trị viên tra cứu tức thì thông tin đơn vé và khách hàng đang ngồi tại một vị trí ghế cụ thể trong phòng chiếu khi có sự cố phát sinh.",
        "steps": [
            "Bước 1: Nhân viên truy cập phân hệ \"Xử lý sự cố\" trên giao diện quản trị/nội bộ.",
            "Bước 2: Nhân viên nhập Mã đặt vé / Số điện thoại của khách hàng hoặc chọn trực tiếp theo Phòng chiếu → Suất chiếu → Vị trí ghế ngồi gặp sự cố.",
            "Bước 3: Hệ thống truy vấn ngược và hiển thị chi tiết đơn đặt vé liên quan: Họ tên khách hàng, Số điện thoại, Mã vé, Loại vé, Giá vé và Tình trạng thanh toán."
        ],
        "notes": [
            "Tính năng này áp dụng phân quyền nghiêm ngặt theo cụm rạp để đảm bảo tính riêng tư của dữ liệu khách hàng."
        ]
    },
    {
        "num": 38,
        "id": "UC-38",
        "name": "Xử lý đổi chỗ ngồi cho khách do lỗi kỹ thuật",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Xử lý nghiệp vụ đổi chỗ ngồi cho khách sang vị trí ghế mới tương đương khi ghế cũ bị hỏng, ướt hoặc gặp lỗi kỹ thuật, đồng thời giữ nguyên mã vé ban đầu.",
        "steps": [
            "Bước 1: Tại màn hình xử lý sự cố, nhân viên chọn đơn vé và vị trí ghế cũ gặp sự cố của khách hàng.",
            "Bước 2: Hệ thống hiển thị sơ đồ phòng chiếu với các vị trí ghế còn trống trong cùng suất chiếu.",
            "Bước 3: Nhân viên chọn vị trí ghế mới phù hợp cho khách hàng.",
            "Bước 4: Nhân viên có thể chọn thêm phương án đền bù thiện chí bằng Voucher quà tặng (Combo bắp nước, Voucher giảm giá) theo quy định chăm sóc khách hàng.",
            "Bước 5: Nhân viên nhấn \"Xác nhận chuyển ghế\".",
            "Bước 6: Hệ thống cập nhật điều chuyển vị trí ghế tại chỗ, giữ nguyên mã vé/mã QR cũ, in lại vé giấy mới với vị trí ghế đã đổi và phát voucher đền bù cho khách (nếu có)."
        ],
        "notes": [
            "Việc đổi ghế trong cùng suất chiếu không phát sinh chênh lệch tiền vé; hệ thống ghi vết đầy đủ nhân viên thao tác và lý do đổi ghế vào lịch sử sự cố."
        ]
    },
    {
        "num": 39,
        "id": "UC-39",
        "name": "Xử lý hủy chỗ và đền bù bằng Voucher quà tặng",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Xử lý tình huống bất khả kháng khi phòng chiếu hết ghế trống để đổi cho khách; hệ thống tiến hành hủy chỗ và đền bù toàn bộ bằng Voucher vé xem phim miễn phí.",
        "steps": [
            "Bước 1: Nhân viên tra cứu đơn vé và chọn chức năng \"Hủy chỗ & Đền bù sự cố\".",
            "Bước 2: Nhân viên chọn vị trí ghế cần hủy và nhập lý do sự cố (ghế hỏng toàn diện, sự cố kỹ thuật phòng chiếu).",
            "Bước 3: Hệ thống tự động tính toán giá trị đền bù tương ứng với 100% giá trị tiền vé của ghế bị hủy.",
            "Bước 4: Nhân viên nhấn nút \"Xác nhận hủy chỗ và đền bù\".",
            "Bước 5: Hệ thống hủy vị trí ghế trong đơn hàng, giải phóng trạng thái ghế và tự động phát một Voucher đền bù 100% giá vé vào ví tài khoản của khách hàng (hoặc gửi quà đền bù tại quầy cho khách vãng lai)."
        ],
        "notes": [
            "Hệ thống không thực hiện hoàn tiền mặt mà giải quyết bồi thường thông qua chính sách phát hành Voucher quà tặng theo quy định vận hành của rạp."
        ]
    },
    {
        "num": 40,
        "id": "UC-40",
        "name": "Khóa bảo trì ghế hỏng vật lý tại phòng chiếu",
        "priority": "Cao",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cho phép nhân viên hoặc quản lý lập tức khóa trạng thái của một chiếc ghế bị gãy, rách hoặc hỏng hóc kỹ thuật sang chế độ bảo trì để ngăn chặn việc bán vé ở mọi suất chiếu tiếp theo.",
        "steps": [
            "Bước 1: Tại màn hình xử lý sự cố hoặc quản lý sơ đồ ghế, nhân viên chọn chiếc ghế bị hỏng.",
            "Bước 2: Nhân viên chọn thao tác \"Khóa bảo trì ghế\".",
            "Bước 3: Nhân viên nhập mô tả tình trạng hư hỏng (ví dụ: gãy tay vịn, hỏng đệm ngồi).",
            "Bước 4: Nhân viên nhấn \"Xác nhận khóa bảo trì\".",
            "Bước 5: Hệ thống chuyển trạng thái vật lý của ghế sang \"Bảo trì (MAINTENANCE)\"; ghế này sẽ lập tức biến mất khỏi danh sách ghế trống của toàn bộ các suất chiếu trong tương lai."
        ],
        "notes": [
            "Khi ghế đã được sửa chữa xong, quản lý có thể thao tác mở lại trạng thái hoạt động bình thường cho ghế."
        ]
    },
    {
        "num": 41,
        "id": "UC-41",
        "name": "Tạo yêu cầu phê duyệt Hủy hóa đơn bắp nước đã bán",
        "priority": "Trung bình",
        "actor": "Nhân viên bán vé & soát vé",
        "description": "Cho phép nhân viên quầy gửi yêu cầu hủy hóa đơn bắp nước lên Quản lý rạp khi khách trả lại đồ hoặc nhân viên nhập nhầm món cần sửa sai.",
        "steps": [
            "Bước 1: Nhân viên tra cứu hóa đơn bán bắp nước cần hủy trong ngày.",
            "Bước 2: Nhân viên nhấn chọn chức năng \"Yêu cầu hủy hóa đơn (Void)\".",
            "Bước 3: Nhân viên nhập lý do yêu cầu hủy (khách đổi ý, pha nhầm món...).",
            "Bước 4: Nhân viên nhấn \"Gửi yêu cầu phê duyệt\".",
            "Bước 5: Hệ thống đưa yêu cầu vào Hàng đợi phê duyệt (Approval Queue) để chờ Quản lý rạp hoặc Quản trị viên xem xét xử lý."
        ],
        "notes": [
            "Nhân viên không thể tự ý hủy hóa đơn đã thanh toán nếu chưa có sự phê duyệt chính thức từ Quản trị viên hoặc Quản lý cơ sở."
        ]
    },

    # =========================================================================
    # KHỐI 3: TÁC NHÂN QUẢN TRỊ VIÊN & QUẢN LÝ (ADMIN / MANAGER) (UC-42 -> UC-72)
    # =========================================================================
    {
        "num": 42,
        "id": "UC-42",
        "name": "Đăng nhập Trang quản trị hệ thống",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Xác thực tài khoản của Quản trị viên cấp cao hoặc Quản lý cụm rạp để cấp quyền điều hành toàn bộ các phân hệ quản trị của DevCine.",
        "steps": [
            "Bước 1: Quản trị viên truy cập đường dẫn cổng quản trị hệ thống.",
            "Bước 2: Nhập tên đăng nhập và mật khẩu tài khoản quản trị.",
            "Bước 3: Nhấn nút \"Đăng nhập\".",
            "Bước 4: Hệ thống xác thực danh tính, kiểm tra quyền hạn chi tiết và tải giao diện Bảng điều khiển quản trị."
        ],
        "notes": [
            "Tài khoản người dùng thông thường nếu cố tình truy cập vào trang quản trị sẽ bị từ chối và tự động chuyển hướng về trang lỗi phân quyền."
        ]
    },
    {
        "num": 43,
        "id": "UC-43",
        "name": "Xem Bảng điều khiển (Dashboard) thống kê Tổng quan Doanh thu, Số vé, Khách mới và Tỷ lệ lấp đầy",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cung cấp bức tranh toàn cảnh về hiệu quả kinh doanh của rạp thông qua các biểu đồ số liệu thời gian thực: Doanh thu, Lượng vé, Khách hàng mới và Tỷ lệ lấp đầy ghế.",
        "steps": [
            "Bước 1: Quản trị viên truy cập mục \"Tổng quan (Dashboard)\".",
            "Bước 2: Chọn khoảng thời gian cần theo dõi: Hôm nay, Tuần này, Tháng này hoặc chọn Tháng/Năm cụ thể.",
            "Bước 3: Hệ thống tổng hợp và hiển thị các thẻ số liệu chính: Tổng doanh thu, Tổng số vé bán ra, Số lượng khách mới, Tỷ lệ lấp đầy phòng chiếu kèm tỷ lệ tăng trưởng so với kỳ trước.",
            "Bước 4: Hệ thống hiển thị Biểu đồ diễn biến doanh thu/số vé theo ngày và Bảng xếp hạng Top các bộ phim ăn khách nhất."
        ],
        "notes": [
            "Quản lý cụm rạp chỉ nhìn thấy số liệu thống kê thuộc cụm rạp mình phụ trách; Quản trị viên cấp cao xem được số liệu toàn hệ thống."
        ]
    },
    {
        "num": 44,
        "id": "UC-44",
        "name": "Quản lý Danh sách Phim",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cho phép quản trị viên thêm mới phim, chỉnh sửa thông tin phim, tải lên hình ảnh áp-phích (poster) và hình nền (banner) qua dịch vụ lưu trữ đám mây.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Quản lý Phim\".",
            "Bước 2: Hệ thống hiển thị danh sách toàn bộ các bộ phim kèm bộ lọc tìm kiếm.",
            "Bước 3: Quản trị viên nhấn \"Thêm phim mới\" hoặc chọn \"Chỉnh sửa\" một bộ phim.",
            "Bước 4: Nhập các trường thông tin: Tên phim, Tên tiếng Anh, Thể loại, Thời lượng, Ngày phát hành, Đạo diễn, Diễn viên, Giới hạn độ tuổi, Mô tả nội dung và Đường dẫn Trailer.",
            "Bước 5: Tải hình ảnh Poster và Banner từ máy tính; hệ thống tự động tải lên dịch vụ đám mây và nhận đường dẫn ảnh.",
            "Bước 6: Nhấn \"Lưu phim\". Hệ thống kiểm tra dữ liệu và lưu thông tin vào cơ sở dữ liệu."
        ],
        "notes": [
            "Thời lượng phim phải là số nguyên dương và các trường thông tin quan trọng không được để trống."
        ]
    },
    {
        "num": 45,
        "id": "UC-45",
        "name": "Quản lý Trạng thái Phim",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Điều chỉnh trạng thái phát hành của bộ phim (Đang chiếu, Sắp chiếu, Ngừng chiếu) hoặc thực hiện xóa mềm bộ phim khỏi hệ thống.",
        "steps": [
            "Bước 1: Tại danh mục quản lý phim, quản trị viên chọn bộ phim cần cập nhật trạng thái.",
            "Bước 2: Thay đổi trạng thái hiển thị của phim (Đang chiếu / Sắp chiếu / Ngừng chiếu) hoặc nhấn biểu tượng Xóa phim.",
            "Bước 3: Hệ thống hiển thị hộp thoại xác nhận hành động.",
            "Bước 4: Quản trị viên xác nhận; hệ thống cập nhật trạng thái mới của phim trong cơ sở dữ liệu."
        ],
        "notes": [
            "Hệ thống không cho phép xóa vĩnh viễn phim đã từng có phát sinh giao dịch đặt vé nhằm bảo toàn tính toàn vẹn dữ liệu lịch sử."
        ]
    },
    {
        "num": 46,
        "id": "UC-46",
        "name": "Quản lý Danh mục Thể loại phim",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý danh sách các thể loại phim trong hệ thống (Hành động, Tình cảm, Kinh dị, Hoạt hình, Viễn tưởng...) phục vụ việc phân loại và tìm kiếm phim.",
        "steps": [
            "Bước 1: Quản trị viên truy cập mục \"Danh mục phim\" và chọn tab \"Thể loại\".",
            "Bước 2: Hệ thống hiển thị danh sách các thể loại hiện có.",
            "Bước 3: Quản trị viên có thể Thêm mới, Chỉnh sửa tên thể loại hoặc Xóa thể loại.",
            "Bước 4: Nhấn \"Lưu\", hệ thống kiểm tra tính hợp lệ và cập nhật dữ liệu."
        ],
        "notes": [
            "Tên thể loại không được chứa ký tự đặc biệt nguy hiểm và không được phép trùng lặp với thể loại đã tồn tại."
        ]
    },
    {
        "num": 47,
        "id": "UC-47",
        "name": "Quản lý Danh mục Định dạng chiếu",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý các định dạng công nghệ trình chiếu phim (2D, 3D, IMAX, 4DX, ScreenX) để gán cho phòng chiếu, suất chiếu và cấu hình phụ thu giá vé.",
        "steps": [
            "Bước 1: Quản trị viên chọn tab \"Định dạng chiếu\" trong phần Danh mục phim.",
            "Bước 2: Hệ thống hiển thị danh sách các định dạng phim đang hỗ trợ.",
            "Bước 3: Quản trị viên thực hiện Thêm mới, Chỉnh sửa mã định dạng/tên định dạng hoặc Xóa định dạng.",
            "Bước 4: Hệ thống lưu và đồng bộ danh mục định dạng trên toàn hệ thống."
        ],
        "notes": [
            "Không thể xóa định dạng nếu đang có phòng chiếu hoặc suất chiếu sử dụng định dạng này."
        ]
    },
    {
        "num": 48,
        "id": "UC-48",
        "name": "Quản lý Danh mục Phân loại độ tuổi kiểm duyệt",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý các nhãn kiểm duyệt độ tuổi khán giả theo quy định của Cục Điện ảnh (P - Phổ biến, K - Dưới 13 tuổi có phụ huynh, T13, T16, T18, C - Cấm phổ biến).",
        "steps": [
            "Bước 1: Quản trị viên chọn tab \"Kiểm duyệt độ tuổi\" trong Danh mục phim.",
            "Bước 2: Hệ thống hiển thị danh sách các mã phân loại độ tuổi và mô tả quy định tương ứng.",
            "Bước 3: Quản trị viên có thể thêm mã mới, chỉnh sửa giải thích độ tuổi hoặc xóa nhãn không còn áp dụng.",
            "Bước 4: Nhấn \"Lưu thay đổi\" để cập nhật vào hệ thống."
        ],
        "notes": [
            "Mã kiểm duyệt được chuẩn hóa tự động viết hoa (ví dụ: T13, T18) để bảo đảm tính thống nhất."
        ]
    },
    {
        "num": 49,
        "id": "UC-49",
        "name": "Quản lý Cụm rạp chi nhánh",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cho phép thêm mới cụm rạp chi nhánh, cập nhật thông tin tên rạp, địa chỉ chi tiết, khu vực tỉnh/thành phố và thông tin liên hệ của từng rạp.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Quản lý Rạp chiếu\".",
            "Bước 2: Hệ thống hiển thị danh sách các cụm rạp hiện có trên toàn quốc.",
            "Bước 3: Quản trị viên nhấn \"Thêm cụm rạp mới\" hoặc chọn rạp để \"Chỉnh sửa\".",
            "Bước 4: Nhập thông tin: Tên cụm rạp, Tỉnh/Thành phố, Địa chỉ chi tiết, Số điện thoại liên hệ và hình ảnh đại diện.",
            "Bước 5: Nhấn \"Lưu cụm rạp\", hệ thống cập nhật vào cơ sở dữ liệu."
        ],
        "notes": [
            "Mỗi cụm rạp là một không gian vận hành độc lập phục vụ cho việc phân quyền nhân sự và cách ly số liệu doanh thu."
        ]
    },
    {
        "num": 50,
        "id": "UC-50",
        "name": "Quản lý Phòng chiếu",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý danh sách các phòng chiếu (halls) thuộc từng cụm rạp, thiết lập định dạng hỗ trợ của phòng và theo dõi tổng sức chứa ghế ngồi.",
        "steps": [
            "Bước 1: Quản trị viên chọn một cụm rạp cụ thể và truy cập tab \"Cơ sở vật chất\".",
            "Bước 2: Hệ thống hiển thị danh sách các phòng chiếu kèm định dạng phòng (2D, 3D, IMAX) và tổng số lượng ghế.",
            "Bước 3: Quản trị viên nhấn \"Thêm phòng chiếu\", nhập tên phòng (ví dụ: Phòng 01, Cinema 2) và chọn định dạng chiếu.",
            "Bước 4: Nhấn \"Xác nhận tạo phòng\". Hệ thống lưu phòng chiếu mới vào cụm rạp."
        ],
        "notes": [
            "Hệ thống yêu cầu xác nhận kỹ lưỡng khi xóa phòng chiếu để tránh ảnh hưởng đến các suất chiếu đang được xếp lịch."
        ]
    },
    {
        "num": 51,
        "id": "UC-51",
        "name": "Thiết kế và Lưu trữ Sơ đồ ma trận ghế phòng chiếu",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cung cấp công cụ vẽ và chỉnh sửa ma trận ghế trực quan: Tạo kích thước hàng/cột, gắn nhãn ghế (A1, B2...), thiết lập loại ghế (Thường, VIP, Đôi, Lối đi).",
        "steps": [
            "Bước 1: Quản trị viên chọn một phòng chiếu và nhấn chức năng \"Sơ đồ ghế\".",
            "Bước 2: Hệ thống mở giao diện công cụ thiết kế ma trận ghế trực quan.",
            "Bước 3: Quản trị viên cấu hình số hàng và số cột tổng thể của phòng.",
            "Bước 4: Sử dụng công cụ cọ vẽ (brush) để phân bổ các vị trí: Ghế thường, Ghế VIP, Ghế Sweetbox hoặc làm Khoảng trống / Lối đi.",
            "Bước 5: Kiểm tra bảng xem trước sơ đồ ghế và nhấn nút \"Lưu sơ đồ ghế\".",
            "Bước 6: Hệ thống tự động khởi tạo danh sách ghế thực tế trong cơ sở dữ liệu theo đúng sơ đồ vừa vẽ."
        ],
        "notes": [
            "Việc sửa đổi sơ đồ ghế chỉ được áp dụng cho các suất chiếu được tạo mới sau thời điểm lưu sơ đồ."
        ]
    },
    {
        "num": 52,
        "id": "UC-52",
        "name": "Lập lịch và Điều phối suất chiếu",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Hỗ trợ xếp lịch chiếu phim trên dòng thời gian trực quan, hỗ trợ tạo suất lẻ hoặc tạo hàng loạt (batch scheduling), tự động kiểm tra xung đột khung giờ chiếu.",
        "steps": [
            "Bước 1: Quản trị viên truy cập tab \"Lịch chiếu\" của cụm rạp và chọn Ngày cần xếp lịch.",
            "Bước 2: Hệ thống hiển thị dòng thời gian trực quan của tất cả các phòng chiếu trong ngày.",
            "Bước 3: Quản trị viên nhấn vào khung giờ trống hoặc nhấn \"Tạo suất chiếu\".",
            "Bước 4: Chọn Phim, Phòng chiếu, Định dạng, Thời gian bắt đầu và Giá vé áp dụng.",
            "Bước 5: Hệ thống tự động cộng thời lượng phim cùng thời gian dọn dẹp vệ sinh phòng (mặc định 15 phút) để tính thời gian kết thúc và kiểm tra xung đột với các suất chiếu khác trong cùng phòng.",
            "Bước 6: Nếu không có xung đột, quản trị viên nhấn \"Xuất bản suất chiếu\" để mở bán."
        ],
        "notes": [
            "Hệ thống lập tức cảnh báo màu đỏ và ngăn chặn lưu nếu phát hiện thời gian hai suất chiếu bị chồng lấn lên nhau."
        ]
    },
    {
        "num": 53,
        "id": "UC-53",
        "name": "Điều chỉnh hoặc Hủy bỏ Suất chiếu",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cho phép thay đổi giờ chiếu, chuyển phòng chiếu hoặc hủy bỏ một suất chiếu đã lên lịch khi có sự cố kỹ thuật hoặc thay đổi kế hoạch chiếu.",
        "steps": [
            "Bước 1: Trên dòng thời gian lịch chiếu, quản trị viên chọn một suất chiếu cụ thể.",
            "Bước 2: Xem chi tiết tình trạng suất: Số vé đã bán, doanh thu tạm tính và danh sách ghế đã được đặt.",
            "Bước 3: Quản trị viên có thể chỉnh sửa khung giờ bắt đầu hoặc nhấn nút \"Hủy suất chiếu\".",
            "Bước 4: Hệ thống kiểm tra: Nếu suất chiếu đã có khách đặt vé, hệ thống yêu cầu xử lý đền bù/chuyển vé trước khi hủy.",
            "Bước 5: Xác nhận thao tác, hệ thống cập nhật trạng thái suất chiếu thành Đã hủy (Cancelled)."
        ],
        "notes": [
            "Hủy suất chiếu đã có vé bán ra là thao tác nghiêm trọng và được ghi vết chi tiết vào Nhật ký kiểm toán hệ thống."
        ]
    },
    {
        "num": 54,
        "id": "UC-54",
        "name": "Quản lý Thực đơn Bắp nước",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý danh sách các món ăn, nước ngọt và các gói combo ẩm thực: Thêm món, chỉnh sửa giá bán, tải ảnh minh họa và bật/tắt trạng thái mở bán.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Thực đơn F&B\".",
            "Bước 2: Hệ thống hiển thị danh sách toàn bộ các sản phẩm ẩm thực đang có.",
            "Bước 3: Nhấn \"Thêm món mới\" hoặc chọn món để \"Chỉnh sửa\".",
            "Bước 4: Nhập Tên sản phẩm, Loại món (Combo, Bắp rang, Nước uống, Snack), Đơn giá niêm yết, Hình ảnh sản phẩm và Mô tả thành phần.",
            "Bước 5: Thiết lập trạng thái Bật/Tắt kinh doanh.",
            "Bước 6: Nhấn \"Lưu thông tin\", hệ thống cập nhật thực đơn tức thì trên toàn hệ thống đặt vé online và POS."
        ],
        "notes": [
            "Hệ thống quản lý tồn kho vô hạn, tập trung hoàn toàn vào việc hiển thị thực đơn và định giá bán chính xác."
        ]
    },
    {
        "num": 55,
        "id": "UC-55",
        "name": "Quản lý Nhóm tùy chọn Bắp nước",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý kho các tùy chọn khẩu vị (Vị phô mai, Vị caramel, Nước suối, Nước ngọt các loại) và gán các ô chọn vị linh hoạt vào từng combo.",
        "steps": [
            "Bước 1: Quản trị viên chọn tab \"Kho tùy chọn\" trong phân hệ Thực đơn F&B.",
            "Bước 2: Quản trị viên khởi tạo các nhóm tùy chọn (ví dụ: Nhóm Vị Bắp, Nhóm Loại Nước) và danh sách các vị kèm mức giá phụ thu thêm (nếu có).",
            "Bước 3: Khi cấu hình Combo, quản trị viên thêm các ô lựa chọn (Slots) và gán nhóm tùy chọn tương ứng (ví dụ: Ô 1 chọn 1 vị bắp, Ô 2 chọn 2 loại nước).",
            "Bước 4: Nhấn \"Lưu cấu hình combo\". Hệ thống tự động sinh giao diện chọn vị thông minh cho khách hàng và thu ngân."
        ],
        "notes": [
            "Hỗ trợ thiết lập số lượng lựa chọn tối thiểu và tối đa trên từng ô để đảm bảo khách chọn đúng số lượng món trong combo."
        ]
    },
    {
        "num": 56,
        "id": "UC-56",
        "name": "Quản lý Chương trình Khuyến mãi",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Khởi tạo và quản lý các chương trình ưu đãi, voucher giảm giá: Thiết lập loại giảm giá (phần trăm hoặc tiền cố định), điều kiện đơn hàng, thời gian áp dụng.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Khuyến mãi & Ưu đãi\".",
            "Bước 2: Hệ thống hiển thị danh sách các chương trình khuyến mãi hiện có.",
            "Bước 3: Quản trị viên nhấn \"Tạo chương trình mới\".",
            "Bước 4: Thiết lập các thông số: Mã code khuyến mãi (ví dụ: SUMMER2026), Tên chương trình, Loại giảm giá (% hoặc Tiền mặt), Giá trị giảm, Giảm tối đa, Đơn hàng tối thiểu, Số vé tối đa áp dụng, Ngày bắt đầu, Ngày kết thúc và Đối tượng khách hàng áp dụng.",
            "Bước 5: Cấu hình tùy chọn cho phép khách hàng đổi bằng điểm thưởng tích lũy (nếu áp dụng).",
            "Bước 6: Nhấn \"Lưu khuyến mãi\". Hệ thống kích hoạt chương trình theo đúng lịch trình."
        ],
        "notes": [
            "Mã khuyến mãi đang trong thời gian chạy sẽ được khóa ngày bắt đầu để đảm bảo tính nhất quán của dữ liệu giao dịch."
        ]
    },
    {
        "num": 57,
        "id": "UC-57",
        "name": "Phát Mã giảm giá trực tiếp cho Khách hàng cụ thể",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cho phép quản trị viên cấp phát trực tiếp một mã giảm giá hoặc phiếu ưu đãi đặc biệt vào ví tài khoản của một khách hàng cụ thể để tri ân hoặc hỗ trợ.",
        "steps": [
            "Bước 1: Tại danh sách chương trình khuyến mãi, quản trị viên chọn thao tác \"Phát voucher cho khách\".",
            "Bước 2: Tìm kiếm và chọn tài khoản khách hàng cần phát mã qua Số điện thoại hoặc Email.",
            "Bước 3: Thiết lập hạn sử dụng riêng cho voucher của khách (nếu cần).",
            "Bước 4: Nhấn \"Xác nhận phát voucher\".",
            "Bước 5: Hệ thống tạo một voucher mới gắn liền với tài khoản của khách hàng và gửi thông báo vào ứng dụng của khách."
        ],
        "notes": [
            "Voucher được phát trực tiếp sẽ hiển thị ngay trong mục \"Ưu đãi của tôi\" trên tài khoản của khách hàng."
        ]
    },
    {
        "num": 58,
        "id": "UC-58",
        "name": "Gửi Chiến dịch Email Thông báo Khuyến mãi cho tập khách hàng đủ điều kiện",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Tự động quét tập khách hàng thỏa mãn tiêu chí của chương trình ưu đãi và gửi thư điện tử quảng bá chiến dịch kèm mã voucher đến hộp thư của khách.",
        "steps": [
            "Bước 1: Quản trị viên chọn một chương trình khuyến mãi đang hoạt động.",
            "Bước 2: Nhấn chọn chức năng \"Gửi email chiến dịch\".",
            "Bước 3: Hệ thống hiển thị số lượng khách hàng thuộc đối tượng áp dụng (Tất cả khách hàng, Khách hàng mới, Khách VIP...) chưa nhận thông báo.",
            "Bước 4: Quản trị viên xác nhận gửi.",
            "Bước 5: Hệ thống kích hoạt tiến trình gửi email hàng loạt kèm mẫu thư thiết kế đẹp mắt và thông báo tổng số thư đã gửi thành công."
        ],
        "notes": [
            "Hệ thống tự động loại trừ các khách hàng đã nhận email mã này trước đó để tránh gửi thư trùng lặp gây phiền hà."
        ]
    },
    {
        "num": 59,
        "id": "UC-59",
        "name": "Thiết lập Quy tắc Bảng giá vé",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cấu hình động công thức tính giá vé xem phim duy nhất cho toàn hệ thống: Giá nền theo Thứ trong tuần × Khung giờ × Đối tượng; Phụ thu loại ghế; Phụ thu định dạng và Phụ thu ngày lễ.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Cấu hình Bảng giá\".",
            "Bước 2: Tại tab \"Giá nền\", thiết lập ma trận mức giá cơ bản theo các ngày trong tuần (Thứ 2 đến Thứ 5, Cuối tuần), các khung giờ (Trước 12h, Sau 12h) và Đối tượng khán giả (Người lớn, Học sinh/Sinh viên).",
            "Bước 3: Tại tab \"Phụ thu loại ghế\", thiết lập mức phụ thu thêm cho Ghế VIP và Ghế đôi Sweetbox.",
            "Bước 4: Tại tab \"Phụ thu định dạng\", thiết lập mức phụ thu thêm cho các định dạng đặc biệt (3D, IMAX).",
            "Bước 5: Tại tab \"Ngày lễ\", quản lý danh sách ngày nghỉ lễ quốc gia và cấu hình mức phụ thu ngày lễ.",
            "Bước 6: Nhấn \"Lưu bảng giá\". Toàn bộ hệ thống đặt vé sẽ áp dụng ngay công thức tính giá mới."
        ],
        "notes": [
            "Công thức tính giá vé máy chủ đảm bảo tính nhất quán tuyệt đối giữa kênh bán trực tuyến (Online) và kênh bán tại quầy (POS)."
        ]
    },
    {
        "num": 60,
        "id": "UC-60",
        "name": "Sử dụng Công cụ Mô phỏng & Tính thử Giá vé trực quan",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cung cấp công cụ mô phỏng trực tiếp giúp quản trị viên kiểm thử nhanh công thức tính giá vé xem phim với các tham số đầu vào giả định trước khi ban hành.",
        "steps": [
            "Bước 1: Quản trị viên truy cập tab \"Tính thử giá vé (Simulator)\" trong phân hệ Bảng giá.",
            "Bước 2: Chọn các tham số thử nghiệm: Ngày chiếu, Giờ chiếu, Loại ghế (Thường/VIP/Đôi), Định dạng phòng chiếu và Đối tượng khán giả.",
            "Bước 3: Hệ thống lập tức tính toán và bóc tách chi tiết cấu thành giá: [Giá nền] + [Phụ thu ghế] + [Phụ thu định dạng] + [Phụ thu ngày lễ] = [Giá vé cuối cùng].",
            "Bước 4: Quản trị viên đối chiếu kết quả để đảm bảo chính sách giá được thiết lập đúng như mong muốn."
        ],
        "notes": [
            "Công cụ chạy thử nghiệm độc lập không làm thay đổi hay ghi dữ liệu giao dịch vào hệ thống."
        ]
    },
    {
        "num": 61,
        "id": "UC-61",
        "name": "Quản lý Danh sách Hóa đơn đặt vé và Xem chi tiết giao dịch toàn hệ thống",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý toàn bộ danh sách các đơn đặt vé trong hệ thống, tìm kiếm lọc theo trạng thái/kênh bán, xem chi tiết vé, bắp nước và thông tin thanh toán đối soát.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Quản lý Hóa đơn (Bookings)\".",
            "Bước 2: Sử dụng bộ lọc tìm kiếm theo: Mã đơn hàng, Trạng thái (Hoàn tất, Đang giữ, Đã hủy), Kênh bán (Online / Quầy POS), Phương thức thanh toán và Khoảng ngày giao dịch.",
            "Bước 3: Nhấn chọn một đơn hàng để mở cửa sổ chi tiết giao dịch.",
            "Bước 4: Hệ thống hiển thị chi tiết: Thông tin khách hàng, Tên phim, Suất chiếu, Chi tiết ghế và giá vé, Chi tiết bắp nước và các tùy chọn vị, Mã giảm giá đã áp dụng, Mã tham chiếu ngân hàng và Trạng thái in vé/check-in."
        ],
        "notes": [
            "Hỗ trợ xem lại bản in biên lai hóa đơn để phục vụ công tác đối chiếu hoặc hỗ trợ giải quyết khiếu nại của khách hàng."
        ]
    },
    {
        "num": 62,
        "id": "UC-62",
        "name": "Quản lý Danh sách Khách hàng và Lịch sử giao dịch",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Theo dõi danh sách người dùng thành viên, cấp bậc thành viên, tổng điểm tích lũy và tổng chi tiêu của từng khách hàng trong toàn hệ thống.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Khách hàng\".",
            "Bước 2: Hệ thống hiển thị danh sách khách hàng kèm Họ tên, Số điện thoại, Email, Hạng thành viên (Đồng, Bạc, Vàng, Bạch Kim) và Điểm thưởng.",
            "Bước 3: Quản trị viên có thể tìm kiếm khách hàng theo tên hoặc số điện thoại.",
            "Bước 4: Nhấn vào một khách hàng để xem lịch sử tất cả các đơn đặt vé mà khách hàng đó đã thực hiện."
        ],
        "notes": [
            "Hệ thống hỗ trợ bảo mật thông tin nhạy cảm của khách hàng theo các tiêu chuẩn an toàn dữ liệu."
        ]
    },
    {
        "num": 63,
        "id": "UC-63",
        "name": "Quản lý Danh sách Nhân viên và Thông tin phân bổ theo Cụm rạp",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý tài khoản nhân viên nội bộ: Khởi tạo tài khoản nhân sự mới, gán vai trò quyền hạn (Quản trị viên, Quản lý, Nhân viên), phân bổ rạp làm việc và khóa tài khoản khi cần.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Quản lý Nhân viên\".",
            "Bước 2: Hệ thống hiển thị danh sách nhân sự kèm vai trò và cụm rạp công tác tương ứng.",
            "Bước 3: Quản trị viên nhấn \"Thêm nhân viên mới\", nhập Họ tên, Tên đăng nhập, Email, Số điện thoại, Vai trò (ADMIN, MANAGER, STAFF) và chọn Cụm rạp phân công.",
            "Bước 4: Nhấn \"Lưu nhân viên\". Hệ thống tạo tài khoản với mật khẩu tạm thời và bật cờ yêu cầu đổi mật khẩu lần đầu.",
            "Bước 5: Quản trị viên có thể thực hiện Khóa / Kích hoạt lại tài khoản hoặc Đặt lại mật khẩu cho nhân viên khi có yêu cầu."
        ],
        "notes": [
            "Nhân viên được gán cụm rạp nào sẽ chỉ có quyền thực hiện bán vé và soát vé trong phạm vi của cụm rạp đó."
        ]
    },
    {
        "num": 64,
        "id": "UC-64",
        "name": "Thiết lập Ma trận Phân quyền người dùng",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cung cấp bảng ma trận phân quyền chi tiết cho phép bật/tắt các quyền hạn Xem, Thêm, Sửa, Xóa trên từng phân hệ chức năng cho từng vai trò người dùng trong hệ thống.",
        "steps": [
            "Bước 1: Quản trị viên cấp cao truy cập mục \"Phân quyền (Roles & Permissions)\".",
            "Bước 2: Hệ thống hiển thị ma trận quyền gồm danh sách các Tính năng (Phim, Rạp, Lịch chiếu, Bán vé POS, Bắp nước, Khuyến mãi, Báo cáo, Nhật ký...) theo các Cột quyền hạn (Xem, Thêm, Sửa, Xóa, Xử lý).",
            "Bước 3: Quản trị viên tích chọn hoặc bỏ chọn các quyền hạn tương ứng cho từng vai trò (Quản lý, Nhân viên).",
            "Bước 4: Nhấn \"Lưu ma trận phân quyền\".",
            "Bước 5: Hệ thống cập nhật và áp dụng quy tắc kiểm tra quyền hạn tức thì trên toàn bộ các yêu cầu thao tác."
        ],
        "notes": [
            "Vai trò Quản trị viên tối cao (Admin) luôn mặc định sở hữu toàn bộ các quyền hạn và không thể bị vô hiệu hóa."
        ]
    },
    {
        "num": 65,
        "id": "UC-65",
        "name": "Quản lý và Phê duyệt Yêu cầu Hủy đơn hàng Bắp nước",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Xem xét danh sách các yêu cầu hủy hóa đơn bắp nước do nhân viên quầy gửi lên và đưa ra quyết định Phê duyệt hoặc Từ chối kèm ghi chú lý do.",
        "steps": [
            "Bước 1: Quản lý hoặc Quản trị viên truy cập mục \"Hàng đợi phê duyệt (Approval Queue)\".",
            "Bước 2: Hệ thống hiển thị danh sách các yêu cầu hủy đơn hàng F&B đang ở trạng thái Chờ duyệt kèm thông tin: Nhân viên tạo yêu cầu, Mã đơn, Giá trị đơn, Thời gian và Lý do hủy.",
            "Bước 3: Quản lý kiểm tra thông tin thực tế và nhấn nút \"Phê duyệt\" hoặc \"Từ chối\".",
            "Bước 4: Nếu từ chối, quản lý nhập ghi chú lý do từ chối.",
            "Bước 5: Hệ thống cập nhật trạng thái đơn hàng, hủy ghi nhận doanh thu nếu được duyệt và thông báo kết quả cho nhân viên quầy."
        ],
        "notes": [
            "Chỉ người dùng có vai trò Quản trị viên hoặc Quản lý cơ sở mới có thẩm quyền thực hiện phê duyệt các yêu cầu này."
        ]
    },
    {
        "num": 66,
        "id": "UC-66",
        "name": "Kiểm duyệt và Quản lý Đánh giá / Bình luận phim của khách hàng",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Giúp quản trị viên kiểm soát nội dung đánh giá của khách hàng: Theo dõi danh sách nhận xét, ẩn các bình luận có nội dung không phù hợp hoặc xóa đánh giá vi phạm.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Quản lý Đánh giá (Reviews)\".",
            "Bước 2: Hệ thống hiển thị danh sách tất cả các đánh giá của người dùng kèm Tên phim, Người đánh giá, Số sao, Nội dung nhận xét, Thời gian và Trạng thái ẩn/hiện.",
            "Bước 3: Quản trị viên có thể nhấn nút \"Ẩn/Hiện\" để tạm thời ẩn bình luận khỏi trang chi tiết phim.",
            "Bước 4: Quản trị viên có thể nhấn \"Xóa\" đối với các đánh giá chứa ngôn từ vi phạm tiêu chuẩn cộng đồng.",
            "Bước 5: Hệ thống cập nhật trạng thái hiển thị và tự động tính toán lại điểm số trung bình của bộ phim."
        ],
        "notes": [
            "Tính năng này bảo vệ hình ảnh thương hiệu và giữ môi trường thảo luận điện ảnh văn minh, lành mạnh."
        ]
    },
    {
        "num": 67,
        "id": "UC-67",
        "name": "Tiếp nhận, Quản lý trạng thái và Phản hồi Yêu cầu Hỗ trợ khách hàng qua Email",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Tiếp nhận các phiếu yêu cầu hỗ trợ từ khách hàng, phân loại, cập nhật tiến độ xử lý và soạn thư phản hồi gửi trực tiếp đến email của khách hàng.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Chăm sóc khách hàng / Hỗ trợ (Customer Support)\".",
            "Bước 2: Hệ thống hiển thị danh sách các phiếu yêu cầu hỗ trợ (Support Tickets) phân loại theo trạng thái (Mới tiếp nhận, Đang xử lý, Đã giải quyết).",
            "Bước 3: Quản trị viên chọn một phiếu hỗ trợ để xem thông tin khách hàng, số điện thoại, chủ đề và nội dung phản ánh.",
            "Bước 4: Quản trị viên nhập nội dung giải đáp/xử lý vào ô \"Phản hồi khách hàng\" và nhấn \"Gửi phản hồi qua Email\".",
            "Bước 5: Hệ thống tự động gửi thư điện tử chứa nội dung trả lời đến email của khách hàng, đồng thời chuyển trạng thái phiếu sang \"Đang xử lý / Đã giải quyết\"."
        ],
        "notes": [
            "Toàn bộ lịch sử trao đổi phản hồi được lưu trữ trong phiếu yêu cầu để tiện theo dõi quá trình hỗ trợ khách hàng."
        ]
    },
    {
        "num": 68,
        "id": "UC-68",
        "name": "Quản lý Danh mục Câu hỏi thường gặp",
        "priority": "Thấp",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cho phép quản trị viên thêm mới, chỉnh sửa nội dung câu hỏi/câu trả lời, sắp xếp thứ tự hiển thị và bật/tắt các mục giải đáp thắc mắc (FAQ) trên website.",
        "steps": [
            "Bước 1: Quản trị viên truy cập mục \"Quản lý FAQ\".",
            "Bước 2: Hệ thống hiển thị danh sách các câu hỏi thường gặp hiện có kèm nhóm danh mục.",
            "Bước 3: Quản trị viên nhấn \"Thêm câu hỏi\", nhập Nhóm danh mục, Tiêu đề câu hỏi, Nội dung giải đáp chi tiết và Thứ tự hiển thị.",
            "Bước 4: Quản trị viên có thể chỉnh sửa nội dung hoặc Bật/Tắt trạng thái hiển thị của từng câu hỏi.",
            "Bước 5: Nhấn \"Lưu\", hệ thống cập nhật tức thì lên trang Hỏi đáp phía khách hàng."
        ],
        "notes": [
            "Các câu hỏi được thiết lập trạng thái ẩn sẽ chỉ hiển thị ở giao diện quản trị và không xuất hiện trên giao diện người dùng."
        ]
    },
    {
        "num": 69,
        "id": "UC-69",
        "name": "Quản lý Đăng tải, Ghim và Điều hướng Banner quảng cáo trang chủ",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý các biểu ngữ hình ảnh (banner/slider) trên trang chủ: Tải ảnh chất lượng cao, ghim vị trí hiển thị ưu tiên và gắn liên kết điều hướng trực tiếp đến phim hoặc sự kiện.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Quản lý Banner\".",
            "Bước 2: Hệ thống hiển thị danh sách các banner quảng cáo đang hoạt động và banner đã lưu trữ.",
            "Bước 3: Quản trị viên nhấn \"Thêm Banner mới\", tải hình ảnh biểu ngữ chất lượng cao lên hệ thống.",
            "Bước 4: Thiết lập Tiêu đề, Phụ đề, Đường dẫn liên kết khi khách bấm vào banner (chuyển đến trang phim, trang khuyến mãi...) và Thứ tự ưu tiên hiển thị.",
            "Bước 5: Thiết lập trạng thái Bật/Tắt hiển thị hoặc Ghim đầu trang.",
            "Bước 6: Nhấn \"Lưu banner\", hệ thống cập nhật biểu ngữ lên màn hình chính của trang chủ."
        ],
        "notes": [
            "Hình ảnh banner được tự động tối ưu hóa kích thước hiển thị để đảm bảo tốc độ tải trang nhanh và đẹp mắt trên mọi thiết bị."
        ]
    },
    {
        "num": 70,
        "id": "UC-70",
        "name": "Quản lý Tin tức và Bài viết sự kiện khuyến mãi",
        "priority": "Trung bình",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Soạn thảo và xuất bản các bài viết tin tức điện ảnh, bài giới thiệu phim hot và bài viết chi tiết về các sự kiện khuyến mãi trên trang web.",
        "steps": [
            "Bước 1: Quản trị viên chọn tab \"Tin bài sự kiện\" trong phân hệ Khuyến mãi.",
            "Bước 2: Hệ thống hiển thị danh sách các bài viết đã xuất bản.",
            "Bước 3: Nhấn \"Tạo bài viết mới\", nhập Tiêu đề bài viết, Mô tả tóm tắt, Ngày bắt đầu sự kiện, Ngày kết thúc sự kiện và tải Ảnh đại diện.",
            "Bước 4: Nhập nội dung bài viết chi tiết và chọn trạng thái Xuất bản (Active) hoặc Lưu nháp (Inactive).",
            "Bước 5: Nhấn \"Lưu bài viết\", hệ thống đăng tải bài viết lên trang tin tức của website."
        ],
        "notes": [
            "Quản trị viên có thể tìm kiếm, chỉnh sửa nội dung hoặc xóa bài viết đã hết hạn sự kiện bất cứ lúc nào."
        ]
    },
    {
        "num": 71,
        "id": "UC-71",
        "name": "Xem Nhật ký hoạt động và Ghi vết thao tác của người dùng trên hệ thống",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Theo dõi và giám sát toàn bộ lịch sử thao tác của nhân viên và quản trị viên (Audit Logs): Đăng nhập, Thêm/Sửa/Xóa dữ liệu, Bán vé, Duyệt hủy đơn nhằm phục vụ công tác an ninh và kiểm toán.",
        "steps": [
            "Bước 1: Quản trị viên cấp cao truy cập phân hệ \"Nhật ký hệ thống (Audit Logs)\".",
            "Bước 2: Hệ thống hiển thị bảng nhật ký ghi vết tự động theo thời gian thực: Thời gian thực hiện, Tài khoản người thao tác, Vai trò, Địa chỉ IP, Hành động (ĐĂNG NHẬP, TẠO MỚI, CHỈNH SỬA, XÓA), Phân hệ bị tác động và Chi tiết thay đổi.",
            "Bước 3: Quản trị viên có thể lọc nhật ký theo Khoảng ngày, Loại hành động hoặc Tên nhân viên cụ thể.",
            "Bước 4: Nhấn vào một dòng nhật ký để xem chi tiết thông tin dữ liệu trước và sau khi thay đổi."
        ],
        "notes": [
            "Nhật ký kiểm toán hệ thống là dữ liệu chỉ đọc (Read-only), không một tài khoản nào có quyền chỉnh sửa hoặc xóa dữ liệu nhật ký để đảm bảo tính minh bạch tuyệt đối."
        ]
    },
    {
        "num": 72,
        "id": "UC-72",
        "name": "Quản lý Cấu hình tham số hệ thống chung",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Thiết lập các tham số vận hành toàn cục của hệ thống: Thời gian giữ ghế tạm thời, thời gian cho phép bán vé trễ, giới hạn số vé tối đa trên một đơn hàng và thông tin tài khoản VietQR nhận tiền.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ \"Cài đặt hệ thống (Settings)\".",
            "Bước 2: Hệ thống hiển thị các nhóm tham số cấu hình chung.",
            "Bước 3: Quản trị viên điều chỉnh các giá trị vận hành: Thời gian giữ ghế trực tuyến (phút), Thời gian cho phép bán vé trễ sau khi phim chiếu (phút), Số lượng vé tối đa được phép đặt trên một lần giao dịch (vé).",
            "Bước 4: Cấu hình thông tin tài khoản nhận thanh toán chuyển khoản tại quầy (Tên ngân hàng, Số tài khoản, Tên chủ tài khoản thụ hưởng phục vụ sinh mã VietQR).",
            "Bước 5: Nhấn nút \"Lưu cấu hình\".",
            "Bước 6: Hệ thống lưu trữ các tham số và áp dụng ngay lập tức cho toàn bộ các giao dịch trên toàn hệ thống."
        ],
        "notes": [
            "Các tham số cấu hình này tác động trực tiếp đến logic vận hành thời gian thực của cả cổng khách hàng và quầy bán vé."
        ]
    }
]

def main():
    print(f"Total Use Cases: {len(use_cases)}")
    doc = docx.Document()

    # Configure Margins (Normal: 1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Add Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title.add_run("PHỤ LỤC A: ĐẶC TẢ USE CASE HỆ THỐNG DEVCINE")
    format_run(r_title, size_pt=16, bold=True, color_rgb=(0, 51, 102))
    title.paragraph_format.space_after = Pt(8)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_desc = desc.add_run("Phụ lục này trình bày chi tiết đặc tả nghiệp vụ (Use Case Specifications) cho toàn bộ 72 chức năng của hệ thống quản lý rạp chiếu phim và đặt vé trực tuyến DevCine, bao gồm luồng thao tác chuẩn của tác nhân, phản hồi của hệ thống và các điều kiện ràng buộc nghiệp vụ liên quan.")
    format_run(r_desc, size_pt=11, italic=True)
    desc.paragraph_format.space_after = Pt(20)

    current_block = ""

    for uc in use_cases:
        # Check block headers
        if uc["num"] == 1:
            h = doc.add_heading(level=1)
            r = h.add_run("A.1. ĐẶC TẢ CÁC USE CASE KHỐI KHÁCH HÀNG (CUSTOMER PORTAL)")
            format_run(r, size_pt=13, bold=True, color_rgb=(0, 51, 102))
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(10)
        elif uc["num"] == 23:
            h = doc.add_heading(level=1)
            r = h.add_run("A.2. ĐẶC TẢ CÁC USE CASE KHỐI NHÂN VIÊN QUẦY & SOÁT VÉ (STAFF / POS / CHECK-IN)")
            format_run(r, size_pt=13, bold=True, color_rgb=(0, 51, 102))
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after = Pt(10)
        elif uc["num"] == 42:
            h = doc.add_heading(level=1)
            r = h.add_run("A.3. ĐẶC TẢ CÁC USE CASE KHỐI QUẢN TRỊ VIÊN & QUẢN LÝ (ADMIN / MANAGER PORTAL)")
            format_run(r, size_pt=13, bold=True, color_rgb=(0, 51, 102))
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after = Pt(10)

        add_use_case_table(
            doc=doc,
            uc_num=uc["num"],
            uc_id=uc["id"],
            uc_name=uc["name"],
            priority=uc["priority"],
            actor=uc["actor"],
            description=uc["description"],
            steps=uc["steps"],
            notes=uc["notes"]
        )

    target_path = "docs/Dac_Ta_Use_Case.docx"
    try:
        doc.save(target_path)
        print(f"Successfully generated {target_path}")
    except PermissionError:
        fallback_path = "docs/Dac_Ta_Use_Case_DevCine.docx"
        doc.save(fallback_path)
        print(f"File '{target_path}' is currently open in Word. Successfully saved to '{fallback_path}'.")

if __name__ == "__main__":
    main()
