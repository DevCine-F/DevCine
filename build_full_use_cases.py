# -*- coding: utf-8 -*-
"""
Script to generate the complete 72 Use Case Specification docx for DevCine Appendix A.
"""
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
        '''
        tblPr[0].append(parse_xml(borders_xml))

def format_run(run, font_name="Times New Roman", size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_use_case_table(doc, uc_num, uc_id, uc_name, priority, actor, description, steps, notes):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="999999", sz="4")

    # Merge cells for rows 2, 3, 4
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(3, 1).merge(table.cell(3, 3))
    table.cell(4, 1).merge(table.cell(4, 3))

    # Row 0: Mã Use case & Tên Use Case
    cell_00 = table.cell(0, 0)
    cell_00.text = "Mã Use case"
    set_cell_background(cell_00, "F2F2F2")
    format_run(cell_00.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_01 = table.cell(0, 1)
    cell_01.text = uc_id
    format_run(cell_01.paragraphs[0].runs[0], bold=True, color_rgb=(0, 51, 102), size_pt=10.5)

    cell_02 = table.cell(0, 2)
    cell_02.text = "Tên Use Case"
    set_cell_background(cell_02, "F2F2F2")
    format_run(cell_02.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_03 = table.cell(0, 3)
    cell_03.text = uc_name
    format_run(cell_03.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    # Row 1: Độ ưu tiên & Tác nhân
    cell_10 = table.cell(1, 0)
    cell_10.text = "Độ ưu tiên"
    set_cell_background(cell_10, "F2F2F2")
    format_run(cell_10.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_11 = table.cell(1, 1)
    cell_11.text = priority
    format_run(cell_11.paragraphs[0].runs[0], size_pt=10.5)

    cell_12 = table.cell(1, 2)
    cell_12.text = "Tác nhân"
    set_cell_background(cell_12, "F2F2F2")
    format_run(cell_12.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_13 = table.cell(1, 3)
    cell_13.text = actor
    format_run(cell_13.paragraphs[0].runs[0], size_pt=10.5)

    # Row 2: Mô tả
    cell_20 = table.cell(2, 0)
    cell_20.text = "Mô tả"
    set_cell_background(cell_20, "F2F2F2")
    format_run(cell_20.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_21 = table.cell(2, 1)
    cell_21.text = description
    format_run(cell_21.paragraphs[0].runs[0], size_pt=10.5)

    # Row 3: Luồng chạy
    cell_30 = table.cell(3, 0)
    cell_30.text = "Luồng chạy"
    set_cell_background(cell_30, "F2F2F2")
    format_run(cell_30.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_31 = table.cell(3, 1)
    cell_31.text = ""
    for idx, step in enumerate(steps):
        p_step = cell_31.paragraphs[0] if idx == 0 else cell_31.add_paragraph()
        p_step.paragraph_format.space_before = Pt(2)
        p_step.paragraph_format.space_after = Pt(2)
        p_step.paragraph_format.line_spacing = 1.15
        run = p_step.add_run(step)
        format_run(run, size_pt=10)

    # Row 4: Lưu ý
    cell_40 = table.cell(4, 0)
    cell_40.text = "Lưu ý"
    set_cell_background(cell_40, "F2F2F2")
    format_run(cell_40.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_41 = table.cell(4, 1)
    cell_41.text = ""
    for idx, note in enumerate(notes):
        p_note = cell_41.paragraphs[0] if idx == 0 else cell_41.add_paragraph()
        p_note.paragraph_format.space_before = Pt(1)
        p_note.paragraph_format.space_after = Pt(1)
        p_note.paragraph_format.line_spacing = 1.15
        run = p_note.add_run(f"- {note}" if not note.startswith("-") else note)
        format_run(run, size_pt=10, italic=True)

    # Apply padding & vertical alignment
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(16)
    r_cap = caption.add_run(f"Bảng A.{uc_num}: Use case {uc_name.lower()}.")
    format_run(r_cap, size_pt=10.5, italic=True)

