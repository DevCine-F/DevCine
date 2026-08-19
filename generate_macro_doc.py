# -*- coding: utf-8 -*-
"""
Script to generate Dac_Ta_Use_Case_Phu_Luc_A_Moi.docx
Containing the detailed 11 Macro/Package Use Case Specifications for DevCine Appendix A.
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
        '''
        tblPr[0].append(parse_xml(borders_xml))

def format_run(run, font_name="Times New Roman", size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_use_case_table(doc, uc_num, uc_id, uc_name, priority, actor, description, steps, notes):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="999999", sz="4")

    # Merge cells for rows 2, 3, 4
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(3, 1).merge(table.cell(3, 3))
    table.cell(4, 1).merge(table.cell(4, 3))

    # Row 0: Mã Use case & Tên Use Case
    cell_00 = table.cell(0, 0)
    cell_00.text = "Mã Use case"
    set_cell_background(cell_00, "F2F2F2")
    format_run(cell_00.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_01 = table.cell(0, 1)
    cell_01.text = uc_id
    format_run(cell_01.paragraphs[0].runs[0], bold=True, color_rgb=(0, 51, 102), size_pt=10.5)

    cell_02 = table.cell(0, 2)
    cell_02.text = "Tên Use Case"
    set_cell_background(cell_02, "F2F2F2")
    format_run(cell_02.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_03 = table.cell(0, 3)
    cell_03.text = uc_name
    format_run(cell_03.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    # Row 1: Độ ưu tiên & Tác nhân
    cell_10 = table.cell(1, 0)
    cell_10.text = "Độ ưu tiên"
    set_cell_background(cell_10, "F2F2F2")
    format_run(cell_10.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_11 = table.cell(1, 1)
    cell_11.text = priority
    format_run(cell_11.paragraphs[0].runs[0], size_pt=10.5)

    cell_12 = table.cell(1, 2)
    cell_12.text = "Tác nhân"
    set_cell_background(cell_12, "F2F2F2")
    format_run(cell_12.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_13 = table.cell(1, 3)
    cell_13.text = actor
    format_run(cell_13.paragraphs[0].runs[0], size_pt=10.5)

    # Row 2: Mô tả
    cell_20 = table.cell(2, 0)
    cell_20.text = "Mô tả"
    set_cell_background(cell_20, "F2F2F2")
    format_run(cell_20.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_21 = table.cell(2, 1)
    cell_21.text = description
    format_run(cell_21.paragraphs[0].runs[0], size_pt=10.5)

    # Row 3: Luồng chạy
    cell_30 = table.cell(3, 0)
    cell_30.text = "Luồng chạy"
    set_cell_background(cell_30, "F2F2F2")
    format_run(cell_30.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_31 = table.cell(3, 1)
    cell_31.text = ""
    for idx, step in enumerate(steps):
        p_step = cell_31.paragraphs[0] if idx == 0 else cell_31.add_paragraph()
        p_step.paragraph_format.space_before = Pt(2)
        p_step.paragraph_format.space_after = Pt(2)
        p_step.paragraph_format.line_spacing = 1.15
        run = p_step.add_run(step)
        format_run(run, size_pt=10)

    # Row 4: Lưu ý
    cell_40 = table.cell(4, 0)
    cell_40.text = "Lưu ý"
    set_cell_background(cell_40, "F2F2F2")
    format_run(cell_40.paragraphs[0].runs[0], bold=True, size_pt=10.5)

    cell_41 = table.cell(4, 1)
    cell_41.text = ""
    for idx, note in enumerate(notes):
        p_note = cell_41.paragraphs[0] if idx == 0 else cell_41.add_paragraph()
        p_note.paragraph_format.space_before = Pt(1)
        p_note.paragraph_format.space_after = Pt(1)
        p_note.paragraph_format.line_spacing = 1.15
        run = p_note.add_run(f"- {note}" if not note.startswith("-") else note)
        format_run(run, size_pt=10, italic=True)

    # Apply padding & vertical alignment
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(16)
    r_cap = caption.add_run(f"Bảng A.{uc_num}: Use case {uc_name.lower()}.")
    format_run(r_cap, size_pt=10.5, italic=True)

macro_use_cases = [
    {
        "num": 1,
        "id": "UC-01",
        "name": "Quản lý tài khoản khách hàng",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Bao gồm toàn bộ quy trình quản lý định danh người dùng: Đăng ký tài khoản thành viên mới, Đăng nhập hệ thống bằng Số điện thoại hoặc Email, Khôi phục mật khẩu qua Email bằng mã xác thực (OTP), Xem và cập nhật thông tin cá nhân, và Đổi mật khẩu tài khoản.",
        "steps": [
            "Bước 1 (Đăng ký): Khách hàng truy cập trang đăng ký, nhập họ tên, số điện thoại, email và mật khẩu bảo mật. Hệ thống kiểm tra dữ liệu hợp lệ và duy nhất, tự động khởi tạo tài khoản hạng Đồng (Bronze) và đăng nhập vào hệ thống.",
            "Bước 2 (Đăng nhập): Khách hàng nhập số điện thoại hoặc email cùng mật khẩu tại form đăng nhập. Hệ thống xác thực danh tính, lưu trạng thái phiên và tải thông tin điểm tích lũy, hạng thẻ.",
            "Bước 3 (Khôi phục mật khẩu): Khi bị quên mật khẩu, khách hàng nhập email để nhận mã xác thực OTP gồm 6 chữ số. Khách hàng nhập đúng mã OTP và thiết lập mật khẩu mới.",
            "Bước 4 (Cập nhật hồ sơ): Khách hàng vào mục Hồ sơ cá nhân để xem thông tin, điểm thưởng, cấp bậc và cập nhật họ tên hoặc thông tin liên lạc khi cần thiết.",
            "Bước 5 (Đổi mật khẩu): Khách hàng nhập mật khẩu hiện tại cùng mật khẩu mới để thay đổi mật khẩu định kỳ nhằm nâng cao tính an toàn."
        ],
        "notes": [
            "- Không cho phép đăng ký hoặc cập nhật số điện thoại/email trùng lặp với tài khoản khác đã có trong hệ thống.",
            "- Mật khẩu phải đáp ứng tiêu chuẩn an toàn từ 8 đến 32 ký tự, bao gồm chữ hoa, chữ thường, số và ký tự đặc biệt.",
            "- Mã xác thực OTP qua email chỉ có hiệu lực ngắn hạn và chỉ được sử dụng một lần duy nhất."
        ]
    },
    {
        "num": 2,
        "id": "UC-02",
        "name": "Tra cứu thông tin phim và cụm rạp",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Hỗ trợ khách hàng tìm kiếm, lọc và xem thông tin chi tiết của các bộ phim (phim đang chiếu, phim sắp chiếu, video trailer, tóm tắt nội dung, diễn viên, giới hạn độ tuổi) và tra cứu danh sách các cụm rạp, phòng chiếu cùng lịch chiếu theo ngày và khu vực.",
        "steps": [
            "Bước 1: Khách hàng truy cập Trang chủ hoặc mục Lịch chiếu, chuyển đổi giữa danh sách \"Phim đang chiếu\" và \"Phim sắp chiếu\".",
            "Bước 2: Khách hàng tìm kiếm theo tên phim hoặc lọc phim theo các tiêu chí: Thể loại (Hành động, Hài, Hoạt hình...), Cụm rạp chiếu, Định dạng (2D, 3D, IMAX).",
            "Bước 3: Khách hàng chọn một bộ phim để xem trang chi tiết: Áp-phích, tên phim, thể loại, thời lượng, phân loại độ tuổi (P, K, T13, T16, T18, C), đạo diễn, diễn viên, tóm tắt nội dung và các suất chiếu theo từng rạp.",
            "Bước 4: Khách hàng nhấn nút \"Xem Trailer\" để mở cửa sổ phát video giới thiệu chính thức của phim.",
            "Bước 5: Khách hàng tra cứu lịch chiếu tổng hợp bằng cách chọn Cụm rạp (theo tỉnh/thành phố) và Ngày chiếu mong muốn để xem danh sách các suất chiếu khả dụng."
        ],
        "notes": [
            "- Chỉ các bộ phim và cụm rạp đang ở trạng thái kích hoạt công khai mới được hiển thị trên giao diện người dùng.",
            "- Các suất chiếu đã diễn ra trong quá khứ sẽ tự động được ẩn đi để tránh nhầm lẫn cho khách hàng.",
            "- Hệ thống hiển thị rõ ràng nhãn cảnh báo giới hạn độ tuổi theo quy định của Cục Điện ảnh."
        ]
    },
    {
        "num": 3,
        "id": "UC-03",
        "name": "Đặt vé trực tuyến và Thanh toán VNPAY (Quy trình cốt lõi)",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Quy trình khép kín cho phép khách hàng chọn suất chiếu, chọn vị trí ghế ngồi trên sơ đồ trực quan (hệ thống tự động khóa ghế tạm thời 10 phút), lựa chọn thêm bắp nước/combo có tùy biến khẩu vị, áp dụng mã giảm giá và thanh toán trực tuyến qua cổng VNPAY.",
        "steps": [
            "Bước 1: Khách hàng chọn một suất chiếu cụ thể và chuyển sang giao diện Sơ đồ phòng chiếu.",
            "Bước 2: Hệ thống hiển thị sơ đồ mặt bằng ghế theo thời gian thực (Ghế thường, Ghế VIP, Ghế Sweetbox). Khách hàng chọn vị trí ghế ngồi và chọn đối tượng khán giả (Người lớn, Học sinh/Sinh viên).",
            "Bước 3: Khách hàng nhấn \"Tiếp tục\", hệ thống thực hiện khóa giữ chỗ tạm thời các ghế đã chọn trong thời gian 10 phút.",
            "Bước 4: Khách hàng chọn thêm các món bắp rang, nước ngọt hoặc combo ưu đãi kèm theo, đồng thời tùy biến lựa chọn vị bắp (phô mai, caramel) và loại nước ngọt tương ứng.",
            "Bước 5: Tại bước thanh toán, khách hàng nhập mã khuyến mãi hoặc chọn voucher từ ví ưu đãi cá nhân để được khấu trừ giảm giá.",
            "Bước 6: Khách hàng kiểm tra tóm tắt đơn hàng, chọn phương thức VNPAY và chuyển hướng sang cổng thanh toán bảo mật VNPAY để quét mã QR hoặc nhập thông tin thẻ.",
            "Bước 7: Sau khi thanh toán thành công, VNPAY phản hồi về hệ thống; hệ thống xác thực chữ ký số, chuyển đơn hàng sang trạng thái Đã xác nhận (Confirmed), hoàn tất đặt chỗ và cộng điểm thưởng thành viên."
        ],
        "notes": [
            "- Hệ thống kiểm tra nghiêm ngặt quy tắc không để lại ghế trống đơn lẻ (orphan seat) ở giữa hoặc đầu hàng ghế.",
            "- Nếu khách hàng không hoàn tất thanh toán trong vòng 10 phút giữ chỗ, hệ thống sẽ tự động hủy đơn và giải phóng ghế về trạng thái trống.",
            "- Số tiền giảm giá được máy chủ tính toán độc lập chống tình trạng gian lận giá vé."
        ]
    },
    {
        "num": 4,
        "id": "UC-04",
        "name": "Quản lý vé điện tử và Lịch sử giao dịch",
        "priority": "Cao",
        "actor": "Khách hàng",
        "description": "Cung cấp vé điện tử (mã đặt vé và mã QR Code) sau khi thanh toán thành công, gửi email tự động xác nhận đơn hàng và cho phép khách hàng tra cứu lại toàn bộ lịch sử các đơn vé đã mua cùng tình trạng sử dụng.",
        "steps": [
            "Bước 1: Ngay sau khi thanh toán thành công, hệ thống hiển thị màn hình vé điện tử chứa đầy đủ thông tin: Mã đặt vé, Tên phim, Cụm rạp, Phòng chiếu, Thời gian chiếu, Vị trí ghế, Danh sách bắp nước và Mã QR Code xác thực.",
            "Bước 2: Hệ thống tự động gửi thư điện tử chứa hóa đơn thanh toán và thông tin vé chi tiết đến địa chỉ email của khách hàng.",
            "Bước 3: Khách hàng đăng nhập và truy cập mục \"Lịch sử đặt vé\" trong trang cá nhân để xem danh sách toàn bộ các đơn hàng theo thứ tự thời gian từ mới nhất đến cũ nhất.",
            "Bước 4: Khách hàng nhấn chọn vào từng đơn hàng để xem lại thông tin chi tiết, tình trạng vé (Đã thanh toán, Đã in vé/Check-in, Đã hủy) và hiển thị lại mã QR để xuất trình cho nhân viên soát vé tại rạp."
        ],
        "notes": [
            "- Dữ liệu lịch sử lưu vết snapshot toàn bộ thông tin giá vé và tên món ăn tại thời điểm giao dịch, bảo toàn tính chính xác vĩnh viễn.",
            "- Khách hàng có thể lưu ảnh mã QR hoặc mở trực tiếp trên website để làm thủ tục vào phòng chiếu."
        ]
    },
    {
        "num": 5,
        "id": "UC-05",
        "name": "Quản lý Ví Voucher và Đánh giá phim",
        "priority": "Trung bình",
        "actor": "Khách hàng",
        "description": "Cung cấp các tính năng tương tác và duy trì khách hàng thân thiết: Lưu trữ và quản lý mã khuyến mãi trong ví cá nhân, sử dụng điểm thưởng tích lũy để đổi voucher giảm giá, và gửi đánh giá chấm điểm sao cho các bộ phim đã xem.",
        "steps": [
            "Bước 1: Khách hàng truy cập mục \"Ưu đãi của tôi\" để xem danh sách voucher khả dụng và lịch sử các voucher đã sử dụng hoặc hết hạn.",
            "Bước 2: Khách hàng nhập mã ưu đãi bí mật nhận được từ các kênh quảng bá để lưu trực tiếp vào ví voucher cá nhân.",
            "Bước 3: Tại tab \"Đổi điểm lấy ưu đãi\", khách hàng sử dụng điểm tích lũy thành viên (Loyalty Points) nhấn \"Đổi ngay\" tại các gói ưu đãi; hệ thống trừ điểm tương ứng và sinh voucher mới vào ví.",
            "Bước 4: Đối với các bộ phim mà khách hàng đã từng mua vé xem thành công, khách hàng truy cập trang chi tiết phim để chấm điểm sao (1 đến 5 sao) và nhập nhận xét cảm nghĩ.",
            "Bước 5: Hệ thống ghi nhận đánh giá, cập nhật điểm trung bình của bộ phim và hiển thị nhận xét công khai trên website."
        ],
        "notes": [
            "- Khách hàng chưa mua vé xem bộ phim đó sẽ không được cấp quyền gửi đánh giá nhằm ngăn ngừa tình trạng đánh giá ảo hoặc tiêu cực vô căn cứ.",
            "- Mỗi gói ưu đãi quy đổi điểm có thể bị giới hạn số lần đổi tối đa cho mỗi tài khoản thành viên."
        ]
    },
    {
        "num": 6,
        "id": "UC-06",
        "name": "Xem Bảng điều khiển (Dashboard) thống kê",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cung cấp bức tranh toàn cảnh về hiệu quả kinh doanh của hệ thống rạp thông qua các biểu đồ số liệu thời gian thực: Tổng doanh thu, Lượng vé bán ra, Số lượng khách hàng mới, Tỷ lệ lấp đầy ghế và bảng xếp hạng các bộ phim ăn khách nhất.",
        "steps": [
            "Bước 1: Quản trị viên đăng nhập trang quản trị và truy cập phân hệ \"Tổng quan (Dashboard)\".",
            "Bước 2: Quản trị viên lựa chọn khoảng thời gian cần theo dõi: Hôm nay, Tuần này, Tháng này hoặc chọn Tháng/Năm cụ thể qua bộ chọn thời gian.",
            "Bước 3: Hệ thống tổng hợp và hiển thị các thẻ chỉ số hiệu suất chính: Tổng doanh thu, Số vé bán ra, Lượng khách hàng mới, Tỷ lệ lấp đầy phòng chiếu kèm tỷ lệ tăng trưởng so với kỳ trước.",
            "Bước 4: Hệ thống hiển thị Biểu đồ diễn biến doanh thu và lượng vé theo ngày, Bảng xếp hạng Top các bộ phim có doanh thu cao nhất và danh sách các giao dịch đặt vé gần nhất."
        ],
        "notes": [
            "- Quản lý chi nhánh chỉ xem được số liệu thống kê thuộc cụm rạp mà mình được phân công phụ trách.",
            "- Quản trị viên cấp cao (Admin) có quyền xem số liệu tổng hợp của toàn hệ thống hoặc lọc theo từng cụm rạp cụ thể."
        ]
    },
    {
        "num": 7,
        "id": "UC-07",
        "name": "Quản lý Hạ tầng Rạp và Sơ đồ ghế",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý danh sách các cụm rạp chi nhánh, thiết lập quy mô phòng chiếu và sử dụng công cụ thiết kế ma trận ghế trực quan (phân bổ ghế Thường, VIP, Đôi Sweetbox, Lối đi và Khóa bảo trì ghế hỏng vật lý).",
        "steps": [
            "Bước 1: Quản trị viên thêm mới hoặc chỉnh sửa thông tin Cụm rạp: Tên cụm rạp, Tỉnh/Thành phố, Địa chỉ chi tiết, Số điện thoại liên hệ và hình ảnh đại diện.",
            "Bước 2: Tại tab Cơ sở vật chất của rạp, quản trị viên thêm mới Phòng chiếu và gán định dạng công nghệ hỗ trợ (2D, 3D, IMAX).",
            "Bước 3: Quản trị viên mở công cụ thiết kế sơ đồ ghế trực quan, thiết lập số hàng và số cột tổng thể của phòng chiếu.",
            "Bước 4: Sử dụng công cụ cọ vẽ (brush) để gán nhãn và phân loại ghế: Ghế thường, Ghế VIP, Ghế đôi Sweetbox hoặc làm Khoảng trống / Lối đi.",
            "Bước 5: Quản trị viên có thể thao tác khóa bảo trì (MAINTENANCE) cho các ghế bị hư hỏng vật lý để chặn không cho bán vé.",
            "Bước 6: Nhấn \"Lưu sơ đồ ghế\", hệ thống tự động khởi tạo danh sách ghế thực tế trong cơ sở dữ liệu."
        ],
        "notes": [
            "- Sơ đồ ghế mới chỉ áp dụng cho các suất chiếu được tạo sau thời điểm lưu sơ đồ.",
            "- Không thể xóa phòng chiếu nếu đang có các suất chiếu được lên lịch hoạt động trong phòng đó."
        ]
    },
    {
        "num": 8,
        "id": "UC-08",
        "name": "Quản lý Phim, Lập lịch và Điều phối suất chiếu",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Kiểm soát dữ liệu phim (thêm phim, cập nhật trạng thái phát hành, tải áp-phích/banner lên máy chủ đám mây, danh mục thể loại, định dạng, độ tuổi) và lập lịch chiếu phim trên dòng thời gian trực quan (tự động kiểm tra xung đột giờ chiếu và thời gian dọn phòng).",
        "steps": [
            "Bước 1: Quản trị viên thêm phim mới, nhập đầy đủ thông tin: Tên phim, Thể loại, Thời lượng, Giới hạn độ tuổi, Ngày khởi chiếu, Đạo diễn, Diễn viên, Tóm tắt nội dung, Trailer và tải ảnh Poster/Banner lên dịch vụ lưu trữ đám mây.",
            "Bước 2: Cập nhật trạng thái hiển thị của phim (Đang chiếu, Sắp chiếu, Ngừng chiếu) hoặc quản lý các danh mục Thể loại, Định dạng, Độ tuổi kiểm duyệt.",
            "Bước 3: Truy cập tab Lịch chiếu của cụm rạp, chọn ngày chiếu và nhấn vào khung giờ trống trên dòng thời gian để tạo suất chiếu lẻ hoặc tạo suất hàng loạt.",
            "Bước 4: Chọn Phim, Phòng chiếu, Định dạng và Thời gian bắt đầu. Hệ thống tự động cộng thời lượng phim cùng thời gian dọn dẹp phòng (15 phút) để tính thời gian kết thúc và kiểm tra xung đột phòng chiếu.",
            "Bước 5: Nếu không có xung đột, quản trị viên nhấn \"Xuất bản suất chiếu\" để mở bán vé trên toàn hệ thống."
        ],
        "notes": [
            "- Hệ thống tự động phát hiện và ngăn chặn hoàn toàn việc lưu nếu thời gian hai suất chiếu trong cùng một phòng bị chồng lấn lên nhau.",
            "- Không cho phép xóa vĩnh viễn phim đã từng phát sinh giao dịch đặt vé nhằm bảo toàn tính toàn vẹn dữ liệu lịch sử."
        ]
    },
    {
        "num": 9,
        "id": "UC-09",
        "name": "Thiết lập Quy tắc Bảng giá vé",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Cấu hình động công thức tính giá vé xem phim áp dụng thống nhất cho toàn hệ thống: Giá nền theo Thứ trong tuần × Khung giờ × Đối tượng; Phụ thu loại ghế; Phụ thu định dạng; Phụ thu ngày lễ và Công cụ mô phỏng tính thử giá vé.",
        "steps": [
            "Bước 1: Quản trị viên truy cập phân hệ Cấu hình Bảng giá.",
            "Bước 2: Thiết lập ma trận Giá nền cơ bản theo các ngày trong tuần (Thứ 2 đến Thứ 5, Cuối tuần), Khung giờ (Trước/Sau 12h) và Đối tượng khán giả (Người lớn, Học sinh/Sinh viên).",
            "Bước 3: Thiết lập mức phụ thu riêng cho Loại ghế (Ghế VIP, Ghế Sweetbox) và Định dạng phòng chiếu (3D, IMAX).",
            "Bước 4: Quản lý danh sách các Ngày lễ quốc gia và cấu hình mức phụ thu áp dụng trong các dịp lễ.",
            "Bước 5: Sử dụng tab \"Tính thử giá vé (Simulator)\" để nhập các điều kiện giả định và kiểm tra công thức bóc tách cấu thành giá vé: [Giá nền] + [Phụ thu ghế] + [Phụ thu định dạng] + [Phụ thu ngày lễ] = [Giá vé cuối cùng].",
            "Bước 6: Nhấn \"Lưu bảng giá\", hệ thống cập nhật và áp dụng công thức mới đồng bộ cho cả kênh đặt vé online và quầy POS."
        ],
        "notes": [
            "- Công thức tính giá vé máy chủ đảm bảo tính nhất quán tuyệt đối giữa kênh bán trực tuyến (Online) và kênh bán tại quầy (POS).",
            "- Công cụ mô phỏng chạy thử nghiệm độc lập không làm thay đổi hay ghi dữ liệu giao dịch vào hệ thống."
        ]
    },
    {
        "num": 10,
        "id": "UC-10",
        "name": "Quản lý Thực đơn F&B và Khuyến mãi",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý kho thực đơn ẩm thực bắp nước (món lẻ, combo, nhóm tùy chọn vị và phụ thu) và triển khai các chương trình khuyến mãi (khởi tạo voucher, thiết lập điều kiện áp dụng, phát mã quà tặng trực tiếp và gửi email chiến dịch hàng loạt).",
        "steps": [
            "Bước 1: Quản trị viên thêm mới sản phẩm ẩm thực (Combo, Bắp rang, Nước uống, Snack), nhập đơn giá, tải ảnh minh họa và thiết lập trạng thái kinh doanh.",
            "Bước 2: Khởi tạo các nhóm tùy chọn (Vị bắp, Loại nước) và gán các ô lựa chọn linh hoạt vào từng combo kèm mức giá phụ thu thêm.",
            "Bước 3: Khởi tạo chương trình Khuyến mãi: Thiết lập Mã code (ví dụ: SUMMER2026), Tên chương trình, Mức giảm giá (% hoặc Tiền mặt), Giảm tối đa, Đơn hàng tối thiểu, Số vé tối đa, Ngày bắt đầu/kết thúc và cờ cho phép đổi bằng điểm thưởng.",
            "Bước 4: Quản trị viên có thể phát mã voucher trực tiếp cho một khách hàng cụ thể hoặc nhấn \"Gửi email chiến dịch\" để hệ thống tự động gửi thư quảng bá kèm voucher đến toàn bộ khách hàng đủ điều kiện."
        ],
        "notes": [
            "- Mã khuyến mãi đang trong thời gian chạy sẽ được khóa ngày bắt đầu để đảm bảo tính nhất quán của dữ liệu giao dịch.",
            "- Hệ thống tự động loại trừ các khách hàng đã nhận email mã này trước đó khi thực hiện gửi chiến dịch hàng loạt."
        ]
    },
    {
        "num": 11,
        "id": "UC-11",
        "name": "Quản lý Hóa đơn, Nhân sự và Phân quyền",
        "priority": "Cao",
        "actor": "Quản trị viên / Quản lý rạp",
        "description": "Quản lý luồng giao dịch hóa đơn bán vé toàn hệ thống, quản lý tài khoản nhân viên phân bổ theo cụm rạp, thiết lập ma trận phân quyền chi tiết (RBAC), kiểm duyệt đánh giá/hỗ trợ khách hàng và giám sát nhật ký kiểm toán hệ thống (Audit Logs).",
        "steps": [
            "Bước 1 (Quản lý Hóa đơn): Quản trị viên tra cứu, tìm kiếm và lọc danh sách đơn đặt vé theo trạng thái, kênh bán, phương thức thanh toán; xem chi tiết vé, bắp nước và mã tham chiếu đối soát ngân hàng.",
            "Bước 2 (Quản lý Nhân sự): Thêm tài khoản nhân viên mới, gán vai trò (ADMIN, MANAGER, STAFF), phân bổ cụm rạp làm việc và kích hoạt cờ yêu cầu đổi mật khẩu lần đầu.",
            "Bước 3 (Thiết lập Phân quyền): Cấu hình bảng ma trận phân quyền (RBAC) chi tiết từng hành động Xem, Thêm, Sửa, Xóa trên từng phân hệ chức năng cho từng vai trò người dùng.",
            "Bước 4 (Vận hành & Phê duyệt): Xem xét và duyệt/từ chối các yêu cầu hủy đơn hàng bắp nước từ quầy; tiếp nhận và phản hồi email phiếu hỗ trợ khách hàng; kiểm duyệt ẩn/xóa bình luận phim vi phạm.",
            "Bước 5 (Nhật ký hệ thống): Xem bảng nhật ký ghi vết tự động (Audit Logs) mọi hành động đăng nhập, thêm, sửa, xóa dữ liệu của người dùng trên toàn hệ thống kèm địa chỉ IP và mốc thời gian thực."
        ],
        "notes": [
            "- Nhân viên bị giới hạn phạm vi thao tác nghiêm ngặt theo cụm rạp trực thuộc (Cinema Scoping); thao tác chéo rạp sẽ bị từ chối truy cập.",
            "- Nhật ký kiểm toán hệ thống là dữ liệu chỉ đọc (Read-only), không một tài khoản nào có quyền chỉnh sửa hoặc xóa dữ liệu nhật ký.",
            "- Vai trò Quản trị viên tối cao (Admin) luôn mặc định sở hữu toàn bộ các quyền hạn và không thể bị vô hiệu hóa."
        ]
    }
]

def main():
    doc = docx.Document()

    # Configure Margins (Normal: 1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title.add_run("PHỤ LỤC A: ĐẶC TẢ USE CASE HỆ THỐNG DEVCINE")
    format_run(r_title, size_pt=16, bold=True, color_rgb=(0, 51, 102))
    title.paragraph_format.space_after = Pt(8)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_desc = desc.add_run("Phụ lục này trình bày chi tiết đặc tả nghiệp vụ (Use Case Specifications) cho các phân hệ chức năng cốt lõi của hệ thống quản lý rạp chiếu phim và đặt vé trực tuyến DevCine, bao gồm luồng thao tác chuẩn của tác nhân, phản hồi của hệ thống và các điều kiện ràng buộc nghiệp vụ liên quan.")
    format_run(r_desc, size_pt=11, italic=True)
    desc.paragraph_format.space_after = Pt(20)

    for uc in macro_use_cases:
        if uc["num"] == 1:
            h = doc.add_heading(level=1)
            r = h.add_run("A.1. ĐẶC TẢ CÁC USE CASE KHỐI KHÁCH HÀNG (CUSTOMER PORTAL)")
            format_run(r, size_pt=13, bold=True, color_rgb=(0, 51, 102))
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(10)
        elif uc["num"] == 6:
            h = doc.add_heading(level=1)
            r = h.add_run("A.2. ĐẶC TẢ CÁC USE CASE KHỐI QUẢN TRỊ VIÊN & QUẢN LÝ (ADMIN / MANAGER PORTAL)")
            format_run(r, size_pt=13, bold=True, color_rgb=(0, 51, 102))
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after = Pt(10)

        add_use_case_table(
            doc=doc,
            uc_num=uc["num"],
            uc_id=uc["id"],
            uc_name=uc["name"],
            priority=uc["priority"],
            actor=uc["actor"],
            description=uc["description"],
            steps=uc["steps"],
            notes=uc["notes"]
        )

    # Save to a new dedicated file, NOT overwriting Dac_Ta_Use_Case.docx
    target_path = "docs/Dac_Ta_Use_Case_Phu_Luc_A_Moi.docx"
    doc.save(target_path)
    print(f"Successfully generated new file: {target_path}")

if __name__ == "__main__":
    main()
