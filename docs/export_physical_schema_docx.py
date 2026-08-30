# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'none')
    tblBorders.append(insideV)
    tblPr.append(tblBorders)

def make_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

def make_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

tables_data = [
    {
        "num": 1,
        "code": "roles",
        "title_vn": "Bảng Vai trò người dùng",
        "rows": [
            [1, "id", "INT", "Khóa chính vai trò", "PK, IDENTITY(1,1)"],
            [2, "name", "VARCHAR(50)", "Tên vai trò (ADMIN, MANAGER, STAFF, CUSTOMER)", "NOT NULL, UNIQUE"],
            [3, "permissions_matrix", "NVARCHAR(MAX)", "Chuỗi JSON mô tả ma trận phân quyền chức năng", "NULL"]
        ]
    },
    {
        "num": 2,
        "code": "users",
        "title_vn": "Bảng Người dùng",
        "rows": [
            [1, "id", "INT", "Khóa chính người dùng", "PK, IDENTITY(1,1)"],
            [2, "username", "VARCHAR(50)", "Tên đăng nhập", "NOT NULL, UNIQUE"],
            [3, "password_hash", "VARCHAR(255)", "Mật khẩu đã băm (BCrypt)", "NOT NULL"],
            [4, "full_name", "NVARCHAR(100)", "Họ và tên hiển thị", "NOT NULL"],
            [5, "avatar_url", "VARCHAR(500)", "Đường dẫn ảnh đại diện (Cloudinary)", "NULL"],
            [6, "email", "VARCHAR(100)", "Địa chỉ email (dùng nhận OTP, thông báo)", "NOT NULL, UNIQUE"],
            [7, "phone", "VARCHAR(20)", "Số điện thoại liên hệ", "NULL"],
            [8, "role_id", "INT", "Vai trò của tài khoản", "FK → roles(id), NOT NULL"],
            [9, "is_active", "BIT", "Trạng thái hoạt động của tài khoản", "NOT NULL, DEFAULT 1"],
            [10, "must_change_password", "BIT", "Bắt buộc đổi mật khẩu ở lần đăng nhập tiếp theo", "NOT NULL, DEFAULT 0"],
            [11, "created_at", "DATETIME", "Thời điểm tạo tài khoản", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 3,
        "code": "customers",
        "title_vn": "Bảng Khách hàng",
        "rows": [
            [1, "user_id", "INT", "Khóa chính hồ sơ khách hàng (quan hệ 1-1 với users)", "PK, FK → users(id)"],
            [2, "dob", "DATE", "Ngày tháng năm sinh", "NULL"],
            [3, "id_card", "VARCHAR(20)", "Số CMND/CCCD", "NULL"],
            [4, "membership_tier", "VARCHAR(20)", "Hạng thành viên", "NOT NULL, DEFAULT 'STANDARD'"],
            [5, "loyalty_points", "INT", "Điểm tích lũy khả dụng hiện tại", "NOT NULL, DEFAULT 0"],
            [6, "lifetime_points", "INT", "Điểm tích lũy trọn đời dùng để xét hạng", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 4,
        "code": "staffs",
        "title_vn": "Bảng Nhân viên",
        "rows": [
            [1, "user_id", "INT", "Khóa chính hồ sơ nhân viên (quan hệ 1-1 với users)", "PK, FK → users(id)"],
            [2, "staff_code", "VARCHAR(20)", "Mã định danh nhân viên", "NOT NULL, UNIQUE"],
            [3, "cinema_id", "INT", "Cụm rạp trực thuộc (phục vụ Cinema Scoping)", "FK → cinemas(id), NOT NULL"],
            [4, "manager_id", "INT", "Quản lý trực tiếp", "FK → staffs(user_id), NULL"],
            [5, "created_at", "DATETIME", "Thời điểm tạo hồ sơ nhân viên", "NOT NULL, DEFAULT GETDATE()"],
            [6, "updated_at", "DATETIME", "Thời điểm cập nhật hồ sơ gần nhất", "NULL"]
        ]
    },
    {
        "num": 5,
        "code": "user_permission_overrides",
        "title_vn": "Bảng Ghi đè quyền người dùng",
        "rows": [
            [1, "id", "INT", "Khóa chính bản ghi ghi đè quyền", "PK, IDENTITY(1,1)"],
            [2, "user_id", "INT", "Người dùng được cấu hình ghi đè", "FK → users(id), NOT NULL"],
            [3, "feature", "VARCHAR(50)", "Mã chức năng hệ thống", "NOT NULL"],
            [4, "action", "VARCHAR(50)", "Hành động trên chức năng", "NOT NULL"],
            [5, "effect", "VARCHAR(10)", "Hiệu lực quyền (ALLOW hoặc DENY)", "NOT NULL"],
            [6, "updated_at", "DATETIME", "Thời điểm cập nhật cấu hình", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 6,
        "code": "audit_logs",
        "title_vn": "Bảng Nhật ký hệ thống",
        "rows": [
            [1, "id", "INT", "Khóa chính bản ghi nhật ký kiểm toán", "PK, IDENTITY(1,1)"],
            [2, "user_id", "INT", "Người dùng thực hiện thao tác", "FK → users(id), NULL"],
            [3, "action", "VARCHAR(50)", "Hành động thực hiện (CREATE, UPDATE, DELETE...)", "NOT NULL"],
            [4, "target_table", "VARCHAR(50)", "Bảng dữ liệu chịu tác động", "NULL"],
            [5, "ip_address", "VARCHAR(50)", "Địa chỉ IP của máy gửi yêu cầu", "NULL"],
            [6, "timestamp", "DATETIME", "Thời điểm ghi nhận thao tác", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 7,
        "code": "notifications",
        "title_vn": "Bảng Thông báo",
        "rows": [
            [1, "id", "INT", "Khóa chính thông báo", "PK, IDENTITY(1,1)"],
            [2, "customer_id", "INT", "Khách hàng nhận thông báo", "FK → customers(user_id), NOT NULL"],
            [3, "title", "NVARCHAR(255)", "Tiêu đề thông báo", "NOT NULL"],
            [4, "message", "NVARCHAR(MAX)", "Nội dung chi tiết thông báo", "NULL"],
            [5, "type", "VARCHAR(50)", "Phân loại thông báo (BOOKING, REMINDER, PROMOTION, SYSTEM)", "NULL"],
            [6, "is_read", "BIT", "Trạng thái đã đọc", "NOT NULL, DEFAULT 0"],
            [7, "created_at", "DATETIME", "Thời điểm tạo thông báo", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 8,
        "code": "support_tickets",
        "title_vn": "Bảng Phiếu hỗ trợ",
        "rows": [
            [1, "id", "INT", "Khóa chính phiếu hỗ trợ", "PK, IDENTITY(1,1)"],
            [2, "customer_id", "INT", "Khách hàng gửi yêu cầu hỗ trợ", "FK → customers(user_id), NOT NULL"],
            [3, "assigned_to_staff", "INT", "Nhân viên tiếp nhận xử lý phiếu", "FK → staffs(user_id), NULL"],
            [4, "issue_type", "NVARCHAR(100)", "Loại vấn đề cần hỗ trợ", "NULL"],
            [5, "description", "NVARCHAR(MAX)", "Mô tả chi tiết vấn đề của khách hàng", "NULL"],
            [6, "phone", "VARCHAR(20)", "Số điện thoại liên hệ lại", "NULL"],
            [7, "status", "VARCHAR(30)", "Trạng thái xử lý phiếu (OPEN, IN_PROGRESS, RESOLVED, CLOSED)", "NOT NULL, DEFAULT 'OPEN'"],
            [8, "admin_reply", "NVARCHAR(MAX)", "Nội dung phản hồi từ nhân viên/quản trị viên", "NULL"],
            [9, "replied_at", "DATETIME", "Thời điểm gửi nội dung phản hồi", "NULL"],
            [10, "created_at", "DATETIME", "Thời điểm tạo phiếu hỗ trợ", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 9,
        "code": "point_transactions",
        "title_vn": "Bảng Giao dịch điểm thưởng",
        "rows": [
            [1, "id", "INT", "Khóa chính giao dịch điểm", "PK, IDENTITY(1,1)"],
            [2, "customer_id", "INT", "Khách hàng phát sinh biến động điểm", "FK → customers(user_id), NOT NULL"],
            [3, "points", "INT", "Số điểm biến động (dương là cộng, âm là trừ)", "NOT NULL"],
            [4, "type", "VARCHAR(30)", "Loại biến động điểm (EARN, REDEEM, ADJUST...)", "NOT NULL"],
            [5, "source", "VARCHAR(50)", "Nguồn phát sinh giao dịch (BOOKING, CONCESSION, PROMOTION...)", "NULL"],
            [6, "ref_code", "VARCHAR(50)", "Mã đơn tham chiếu (booking_code / sale_code)", "NULL"],
            [7, "balance_after", "INT", "Số dư điểm sau khi giao dịch hoàn tất", "NOT NULL"],
            [8, "note", "NVARCHAR(500)", "Ghi chú lý do biến động điểm", "NULL"],
            [9, "created_at", "DATETIME", "Thời điểm phát sinh giao dịch", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 10,
        "code": "cinemas",
        "title_vn": "Bảng Cụm rạp",
        "rows": [
            [1, "id", "INT", "Khóa chính cụm rạp", "PK, IDENTITY(1,1)"],
            [2, "name", "NVARCHAR(100)", "Tên cụm rạp", "NOT NULL"],
            [3, "address", "NVARCHAR(255)", "Địa chỉ chi tiết cụm rạp", "NULL"],
            [4, "city", "NVARCHAR(50)", "Tỉnh/Thành phố", "NULL"],
            [5, "district", "NVARCHAR(50)", "Quận/Huyện", "NULL"],
            [6, "type", "VARCHAR(50)", "Phân loại rạp", "NULL"],
            [7, "hotline", "VARCHAR(20)", "Số điện thoại hotline", "NULL"],
            [8, "rooms", "INT", "Tổng số phòng chiếu", "NOT NULL, DEFAULT 0"],
            [9, "image_url", "VARCHAR(500)", "Đường dẫn ảnh đại diện cụm rạp", "NULL"],
            [10, "description", "NVARCHAR(MAX)", "Mô tả thông tin cụm rạp", "NULL"],
            [11, "latitude", "FLOAT", "Tọa độ vĩ độ định vị bản đồ", "NULL"],
            [12, "longitude", "FLOAT", "Tọa độ kinh độ định vị bản đồ", "NULL"],
            [13, "amenities", "NVARCHAR(500)", "Danh sách tiện ích (phân tách bởi dấu phẩy)", "NULL"],
            [14, "status", "VARCHAR(30)", "Trạng thái hoạt động (ACTIVE, MAINTENANCE, CLOSED)", "NOT NULL, DEFAULT 'ACTIVE'"],
            [15, "opening_time", "TIME", "Giờ mở cửa hàng ngày", "NULL"],
            [16, "closing_time", "TIME", "Giờ đóng cửa hàng ngày", "NULL"],
            [17, "manager_id", "INT", "Quản lý trưởng của cụm rạp", "FK → staffs(user_id), NULL"]
        ]
    },
    {
        "num": 11,
        "code": "rooms",
        "title_vn": "Bảng Phòng chiếu",
        "rows": [
            [1, "id", "INT", "Khóa chính phòng chiếu", "PK, IDENTITY(1,1)"],
            [2, "cinema_id", "INT", "Cụm rạp chứa phòng chiếu", "FK → cinemas(id), NOT NULL"],
            [3, "name", "NVARCHAR(50)", "Tên phòng chiếu", "NOT NULL"],
            [4, "type", "VARCHAR(50)", "Loại phòng chiếu (STANDARD, SUPERPLEX, CINE_COMFORT)", "NOT NULL, DEFAULT 'STANDARD'"],
            [5, "status", "VARCHAR(30)", "Trạng thái hoạt động của phòng", "NOT NULL, DEFAULT 'ACTIVE'"],
            [6, "turnaround_time_mins", "INT", "Thời gian dọn dẹp phòng giữa 2 suất (phút)", "NOT NULL, DEFAULT 15"],
            [7, "matrix_row", "INT", "Số hàng của khung ma trận sơ đồ ghế", "NOT NULL, DEFAULT 0"],
            [8, "matrix_col", "INT", "Số cột của khung ma trận sơ đồ ghế", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 12,
        "code": "seats",
        "title_vn": "Bảng Ghế phòng chiếu",
        "rows": [
            [1, "id", "INT", "Khóa chính ghế", "PK, IDENTITY(1,1)"],
            [2, "room_id", "INT", "Phòng chiếu chứa ghế", "FK → rooms(id), NOT NULL"],
            [3, "row_char", "VARCHAR(5)", "Ký tự đại diện hàng (A, B, C...)", "NOT NULL"],
            [4, "col_num", "INT", "Số thứ tự cột ghế", "NOT NULL"],
            [5, "seat_type_id", "INT", "Loại ghế", "FK → seat_types(id), NOT NULL"],
            [6, "is_active", "BIT", "Ghế còn sử dụng được hay không", "NOT NULL, DEFAULT 1"],
            [7, "label", "VARCHAR(10)", "Nhãn ghế hiển thị (A1, B2...)", "NOT NULL"],
            [8, "custom_label", "BIT", "Cờ đánh dấu nhãn ghế được sửa thủ công", "NOT NULL, DEFAULT 0"],
            [9, "seat_status", "VARCHAR(30)", "Tình trạng ghế (AVAILABLE, MAINTENANCE)", "NOT NULL, DEFAULT 'AVAILABLE'"],
            [10, "grid_row", "INT", "Tọa độ hàng trong lưới sơ đồ vẽ", "NOT NULL"],
            [11, "grid_col", "INT", "Tọa độ cột trong lưới sơ đồ vẽ", "NOT NULL"],
            [12, "cell_kind", "VARCHAR(20)", "Phân loại ô lưới (SEAT: ghế thật, AISLE: lối đi)", "NOT NULL, DEFAULT 'SEAT'"]
        ]
    },
    {
        "num": 13,
        "code": "seat_types",
        "title_vn": "Bảng Loại ghế",
        "rows": [
            [1, "id", "INT", "Khóa chính loại ghế", "PK, IDENTITY(1,1)"],
            [2, "name", "NVARCHAR(50)", "Tên loại ghế (Thường, VIP, Đôi...)", "NOT NULL, UNIQUE"],
            [3, "color_code", "VARCHAR(20)", "Mã màu hiển thị loại ghế trên sơ đồ", "NULL"]
        ]
    },
    {
        "num": 14,
        "code": "movies",
        "title_vn": "Bảng Phim",
        "rows": [
            [1, "id", "INT", "Khóa chính phim", "PK, IDENTITY(1,1)"],
            [2, "title", "NVARCHAR(255)", "Tên gốc của phim", "NOT NULL"],
            [3, "slug", "VARCHAR(255)", "Đường dẫn thân thiện URL", "NOT NULL, UNIQUE"],
            [4, "duration_mins", "INT", "Thời lượng phim (phút)", "NOT NULL"],
            [5, "age_rating", "VARCHAR(20)", "Mã phân loại độ tuổi xem phim (P, K, T13, T16, T18...)", "NULL"],
            [6, "release_date", "DATE", "Ngày phát hành chính thức", "NULL"],
            [7, "start_date", "DATE", "Ngày bắt đầu chiếu tại rạp", "NULL"],
            [8, "end_date", "DATE", "Ngày kết thúc chiếu tại rạp", "NULL"],
            [9, "status", "VARCHAR(30)", "Trạng thái phim (UPCOMING, NOW_SHOWING, ENDED)", "NOT NULL, DEFAULT 'UPCOMING'"],
            [10, "country", "NVARCHAR(100)", "Quốc gia sản xuất", "NULL"],
            [11, "rating", "VARCHAR(10)", "Điểm đánh giá trung bình hiển thị", "NULL"],
            [12, "rating_count", "INT", "Tổng số lượt đánh giá phim", "NOT NULL, DEFAULT 0"],
            [13, "poster_base64", "NVARCHAR(MAX)", "Dữ liệu/Link ảnh poster phim", "NULL"],
            [14, "banner_base64", "NVARCHAR(MAX)", "Dữ liệu/Link ảnh banner phim", "NULL"],
            [15, "show_on_banner", "BIT", "Hiển thị nổi bật trên banner trang chủ", "NOT NULL, DEFAULT 0"],
            [16, "trailer_url", "VARCHAR(500)", "Đường dẫn video trailer (YouTube)", "NULL"],
            [17, "format", "VARCHAR(50)", "Chuỗi mô tả định dạng chiếu", "NULL"],
            [18, "supported_formats", "VARCHAR(255)", "Danh sách các định dạng hỗ trợ", "NULL"],
            [19, "title_vietnamese", "NVARCHAR(255)", "Tên tiếng Việt của phim", "NULL"],
            [20, "production_year", "INT", "Năm sản xuất phim", "NULL"],
            [21, "language", "NVARCHAR(100)", "Ngôn ngữ phim", "NULL"],
            [22, "original_language", "NVARCHAR(100)", "Ngôn ngữ gốc của phim", "NULL"],
            [23, "version_type", "NVARCHAR(50)", "Loại bản chiếu (Lồng tiếng, Phụ đề...)", "NULL"],
            [24, "distributor", "NVARCHAR(150)", "Nhà phát hành phim", "NULL"],
            [25, "director", "NVARCHAR(150)", "Đạo diễn phim", "NULL"],
            [26, "cast_members", "NVARCHAR(MAX)", "Danh sách diễn viên tham gia", "NULL"],
            [27, "description", "NVARCHAR(MAX)", "Tóm tắt nội dung phim", "NULL"],
            [28, "internal_notes", "NVARCHAR(MAX)", "Ghi chú quản lý nội bộ", "NULL"]
        ]
    },
    {
        "num": 15,
        "code": "movie_formats",
        "title_vn": "Bảng Định dạng phim",
        "rows": [
            [1, "id", "INT", "Khóa chính định dạng phim", "PK, IDENTITY(1,1)"],
            [2, "name", "VARCHAR(50)", "Tên định dạng chiếu (2D, 3D, IMAX, 4DX...)", "NOT NULL, UNIQUE"],
            [3, "description", "NVARCHAR(255)", "Mô tả định dạng chiếu", "NULL"],
            [4, "surcharge", "DECIMAL(18,0)", "Phụ thu định dạng ngày thường", "NOT NULL, DEFAULT 0"],
            [5, "weekend_surcharge", "DECIMAL(18,0)", "Phụ thu định dạng cuối tuần và ngày lễ", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 16,
        "code": "categories",
        "title_vn": "Bảng Thể loại phim",
        "rows": [
            [1, "id", "INT", "Khóa chính thể loại phim", "PK, IDENTITY(1,1)"],
            [2, "name", "NVARCHAR(100)", "Tên thể loại phim (Hành động, Hài kịch, Kinh dị...)", "NOT NULL, UNIQUE"],
            [3, "description", "NVARCHAR(255)", "Mô tả thể loại phim", "NULL"]
        ]
    },
    {
        "num": 17,
        "code": "movie_genre_mapping",
        "title_vn": "Bảng Nối phim và thể loại",
        "rows": [
            [1, "movie_id", "INT", "Khóa ngoại trỏ đến phim", "FK → movies(id), NOT NULL"],
            [2, "category_id", "INT", "Khóa ngoại trỏ đến thể loại", "FK → categories(id), NOT NULL"]
        ]
    },
    {
        "num": 18,
        "code": "movie_categories",
        "title_vn": "Bảng Danh mục phim",
        "rows": [
            [1, "movie_id", "INT", "Khóa chính thành phần trỏ đến phim", "PK, FK → movies(id)"],
            [2, "category_id", "INT", "Khóa chính thành phần trỏ đến thể loại", "PK, FK → categories(id)"]
        ]
    },
    {
        "num": 19,
        "code": "age_ratings",
        "title_vn": "Bảng Phân loại độ tuổi",
        "rows": [
            [1, "id", "INT", "Khóa chính độ tuổi", "PK, IDENTITY(1,1)"],
            [2, "code", "VARCHAR(20)", "Mã phân loại độ tuổi (P, K, T13, T16, T18, C)", "NOT NULL, UNIQUE"],
            [3, "name", "NVARCHAR(100)", "Tên hiển thị độ tuổi", "NOT NULL"],
            [4, "description", "NVARCHAR(255)", "Diễn giải chi tiết quy định độ tuổi", "NULL"]
        ]
    },
    {
        "num": 20,
        "code": "showtimes",
        "title_vn": "Bảng Suất chiếu",
        "rows": [
            [1, "id", "INT", "Khóa chính suất chiếu", "PK, IDENTITY(1,1)"],
            [2, "movie_id", "INT", "Phim được chiếu trong suất", "FK → movies(id), NOT NULL"],
            [3, "room_id", "INT", "Phòng chiếu diễn ra suất chiếu", "FK → rooms(id), NOT NULL"],
            [4, "format_id", "INT", "Định dạng chiếu áp dụng", "FK → movie_formats(id), NOT NULL"],
            [5, "start_time", "DATETIME", "Thời điểm bắt đầu chiếu", "NOT NULL"],
            [6, "end_time", "DATETIME", "Thời điểm kết thúc suất chiếu", "NOT NULL"],
            [7, "status", "VARCHAR(30)", "Trạng thái suất chiếu (OPEN, CLOSED, CANCELLED)", "NOT NULL, DEFAULT 'OPEN'"],
            [8, "layout_data", "NVARCHAR(MAX)", "Dữ liệu snapshot sơ đồ ghế tại thời điểm mở suất", "NULL"]
        ]
    },
    {
        "num": 21,
        "code": "bookings",
        "title_vn": "Bảng Đơn đặt vé",
        "rows": [
            [1, "id", "INT", "Khóa chính đơn đặt vé", "PK, IDENTITY(1,1)"],
            [2, "customer_id", "INT", "Khách hàng đặt vé (NULL nếu khách vãng lai mua tại quầy)", "FK → customers(user_id), NULL"],
            [3, "showtime_id", "INT", "Suất chiếu được đặt", "FK → showtimes(id), NOT NULL"],
            [4, "voucher_id", "INT", "Voucher giảm giá áp dụng cho đơn", "FK → vouchers(id), NULL"],
            [5, "sold_by", "INT", "Nhân viên bán vé tại quầy (ghi nhận dấu vết POS)", "FK → staffs(user_id), NULL"],
            [6, "printed_by", "INT", "Nhân viên in vé giấy cho đơn hàng", "FK → staffs(user_id), NULL"],
            [7, "total_price", "DECIMAL(18,0)", "Tổng tiền đơn hàng trước khi giảm giá", "NOT NULL, DEFAULT 0"],
            [8, "final_price", "DECIMAL(18,0)", "Tổng tiền thanh toán thực tế sau giảm giá", "NOT NULL, DEFAULT 0"],
            [9, "payment_method", "VARCHAR(30)", "Phương thức thanh toán (VNPAY, CASH, CARD, TRANSFER)", "NULL"],
            [10, "payment_gateway_ref", "VARCHAR(100)", "Mã tham chiếu giao dịch cổng thanh toán", "NULL"],
            [11, "status", "VARCHAR(30)", "Trạng thái đơn đặt vé (PENDING, CONFIRMED, CANCELLED, EXPIRED)", "NOT NULL, DEFAULT 'PENDING'"],
            [12, "booking_code", "VARCHAR(50)", "Mã đơn đặt vé dùng tra cứu và tạo QR code", "NOT NULL, UNIQUE"],
            [13, "channel", "VARCHAR(20)", "Kênh đặt vé (ONLINE, POS)", "NOT NULL, DEFAULT 'ONLINE'"],
            [14, "pos_terminal_id", "VARCHAR(50)", "Định danh thiết bị POS tạo đơn", "NULL"],
            [15, "created_at", "DATETIME", "Thời điểm tạo đơn đặt vé", "NOT NULL, DEFAULT GETDATE()"],
            [16, "printed_at", "DATETIME", "Thời điểm quét QR và in vé giấy tại quầy", "NULL"],
            [17, "expires_at", "DATETIME", "Hạn chót giữ chỗ chờ thanh toán", "NULL"]
        ]
    },
    {
        "num": 22,
        "code": "booking_seats",
        "title_vn": "Bảng Chi tiết ghế đặt",
        "rows": [
            [1, "id", "INT", "Khóa chính dòng ghế đặt", "PK, IDENTITY(1,1)"],
            [2, "booking_id", "INT", "Đơn đặt vé chứa ghế", "FK → bookings(id), NOT NULL"],
            [3, "seat_id", "INT", "Ghế vật lý được đặt (bị repoint khi đổi ghế)", "FK → seats(id), NOT NULL"],
            [4, "price_snapshot", "DECIMAL(18,0)", "Giá vé chốt tại thời điểm đặt ghế", "NOT NULL, DEFAULT 0"],
            [5, "ticket_type", "VARCHAR(30)", "Đối tượng hưởng giá vé (ADULT, STUDENT, CHILD, SENIOR)", "NOT NULL, DEFAULT 'ADULT'"],
            [6, "status", "VARCHAR(30)", "Trạng thái ghế trong đơn (BOOKED, CANCELLED, CHANGED)", "NOT NULL, DEFAULT 'BOOKED'"]
        ]
    },
    {
        "num": 23,
        "code": "tickets",
        "title_vn": "Bảng Vé điện tử",
        "rows": [
            [1, "id", "INT", "Khóa chính vé điện tử", "PK, IDENTITY(1,1)"],
            [2, "booking_seat_id", "INT", "Dòng ghế đặt tương ứng với vé", "FK → booking_seats(id), NOT NULL, UNIQUE"],
            [3, "qr_code", "VARCHAR(255)", "Mã QR dùng để soát vé vào rạp", "NULL"],
            [4, "is_checked_in", "BIT", "Trạng thái khách đã qua cửa soát vé", "NOT NULL, DEFAULT 0"],
            [5, "is_age_verified", "BIT", "Đã xác thực độ tuổi khách hàng tại cửa kiểm soát", "NOT NULL, DEFAULT 0"],
            [6, "checked_in_by", "INT", "Nhân viên thực hiện soát vé", "FK → staffs(user_id), NULL"],
            [7, "check_in_time", "DATETIME", "Thời điểm thực hiện soát vé", "NULL"]
        ]
    },
    {
        "num": 24,
        "code": "reviews",
        "title_vn": "Bảng Đánh giá phim",
        "rows": [
            [1, "id", "INT", "Khóa chính đánh giá", "PK, IDENTITY(1,1)"],
            [2, "customer_id", "INT", "Khách hàng viết đánh giá", "FK → customers(user_id), NOT NULL"],
            [3, "movie_id", "INT", "Phim được đánh giá", "FK → movies(id), NOT NULL"],
            [4, "ticket_id", "INT", "Vé xem phim dùng xác thực đã xem", "FK → tickets(id), NULL"],
            [5, "rating", "INT", "Điểm đánh giá xếp hạng (1 đến 5 sao)", "NOT NULL"],
            [6, "comment", "NVARCHAR(MAX)", "Nội dung bình luận chi tiết", "NULL"],
            [7, "hidden", "BIT", "Ẩn bình luận khỏi trang công khai (kiểm duyệt)", "NOT NULL, DEFAULT 0"],
            [8, "created_at", "DATETIME", "Thời điểm đăng đánh giá", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },

    {
        "num": 26,
        "code": "fnb_items",
        "title_vn": "Bảng Mặt hàng ẩm thực",
        "rows": [
            [1, "id", "INT", "Khóa chính mặt hàng ẩm thực", "PK, IDENTITY(1,1)"],
            [2, "name", "NVARCHAR(100)", "Tên món bắp nước / combo", "NOT NULL"],
            [3, "type", "VARCHAR(30)", "Phân loại mặt hàng (COMBO, SINGLE)", "NULL"],
            [4, "price", "DECIMAL(18,0)", "Giá bán cơ bản", "NOT NULL, DEFAULT 0"],
            [5, "image_url", "VARCHAR(500)", "Đường dẫn ảnh sản phẩm", "NULL"],
            [6, "description", "NVARCHAR(500)", "Mô tả chi tiết món ẩm thực", "NULL"],
            [7, "is_active", "BIT", "Trạng thái hiển thị bán cho khách", "NOT NULL, DEFAULT 1"],
            [8, "is_deleted", "BIT", "Cờ xóa mềm lưu vết hóa đơn", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 27,
        "code": "fnb_option_groups",
        "title_vn": "Bảng Nhóm tùy chọn ẩm thực",
        "rows": [
            [1, "id", "INT", "Khóa chính nhóm tùy chọn (Option Pool)", "PK, IDENTITY(1,1)"],
            [2, "name", "NVARCHAR(100)", "Tên nhóm tùy chọn (Nước ngọt, Size bắp, Vị bắp...)", "NOT NULL, UNIQUE"]
        ]
    },
    {
        "num": 28,
        "code": "fnb_option_items",
        "title_vn": "Bảng Tùy chọn ẩm thực chi tiết",
        "rows": [
            [1, "id", "INT", "Khóa chính tùy chọn chi tiết", "PK, IDENTITY(1,1)"],
            [2, "group_id", "INT", "Nhóm tùy chọn chứa mặt hàng", "FK → fnb_option_groups(id), NOT NULL"],
            [3, "name", "NVARCHAR(100)", "Tên tùy chọn chi tiết (Coca-Cola, Vị phô mai...)", "NOT NULL"],
            [4, "surcharge_price", "DECIMAL(18,0)", "Phụ thu khi khách chọn tùy chọn này", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 29,
        "code": "fnb_item_slots",
        "title_vn": "Bảng Khe lựa chọn combo ẩm thực",
        "rows": [
            [1, "id", "INT", "Khóa chính khe lựa chọn của combo", "PK, IDENTITY(1,1)"],
            [2, "fnb_item_id", "INT", "Mặt hàng combo chứa khe lựa chọn", "FK → fnb_items(id), NOT NULL"],
            [3, "option_group_id", "INT", "Nhóm tùy chọn liên kết cho khe", "FK → fnb_option_groups(id), NOT NULL"],
            [4, "default_option_item_id", "INT", "Tùy chọn mặc định được gán cho khe", "FK → fnb_option_items(id), NULL"],
            [5, "slot_label", "NVARCHAR(100)", "Nhãn hiển thị khe lựa chọn (Chọn nước 1, Vị bắp...)", "NOT NULL"],
            [6, "display_order", "INT", "Thứ tự hiển thị khe trong combo", "NOT NULL, DEFAULT 0"],
            [7, "min_choices", "INT", "Số lựa chọn tối thiểu bắt buộc", "NOT NULL, DEFAULT 1"],
            [8, "max_choices", "INT", "Số lựa chọn tối đa cho phép", "NOT NULL, DEFAULT 1"],
            [9, "is_required", "BIT", "Cờ bắt buộc phải chọn", "NOT NULL, DEFAULT 1"]
        ]
    },
    {
        "num": 30,
        "code": "booking_fnbs",
        "title_vn": "Bảng Ẩm thực theo đơn đặt vé",
        "rows": [
            [1, "id", "INT", "Khóa chính dòng ẩm thực đặt kèm vé", "PK, IDENTITY(1,1)"],
            [2, "booking_id", "INT", "Đơn đặt vé chứa món ẩm thực", "FK → bookings(id), NOT NULL"],
            [3, "fnb_item_id", "INT", "Mặt hàng ẩm thực được mua", "FK → fnb_items(id), NOT NULL"],
            [4, "item_name_snapshot", "NVARCHAR(100)", "Tên sản phẩm chốt tại thời điểm mua", "NULL"],
            [5, "quantity", "INT", "Số lượng mua", "NOT NULL, DEFAULT 1"],
            [6, "price_snapshot", "DECIMAL(18,0)", "Đơn giá chốt tại thời điểm mua", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 31,
        "code": "booking_fnb_options",
        "title_vn": "Bảng Tùy chọn ẩm thực theo đơn vé",
        "rows": [
            [1, "id", "INT", "Khóa chính tùy chọn ẩm thực đã chọn", "PK, IDENTITY(1,1)"],
            [2, "booking_fnb_id", "INT", "Dòng ẩm thực trong đơn vé", "FK → booking_fnbs(id), NOT NULL"],
            [3, "option_group_id", "INT", "Nhóm tùy chọn của khe đã chọn", "FK → fnb_option_groups(id), NULL"],
            [4, "option_item_id", "INT", "Tùy chọn cụ thể đã chọn", "FK → fnb_option_items(id), NULL"],
            [5, "slot_label_snapshot", "NVARCHAR(100)", "Nhãn khe lựa chọn chốt lúc mua", "NULL"],
            [6, "option_name_snapshot", "NVARCHAR(100)", "Tên tùy chọn chi tiết chốt lúc mua", "NULL"],
            [7, "surcharge_snapshot", "DECIMAL(18,0)", "Giá phụ thu chốt lúc mua", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 32,
        "code": "concession_sales",
        "title_vn": "Bảng Đơn bán bắp nước tại quầy",
        "rows": [
            [1, "id", "INT", "Khóa chính đơn bán ẩm thực tại quầy", "PK, IDENTITY(1,1)"],
            [2, "sale_code", "VARCHAR(50)", "Mã đơn bán ẩm thực tại quầy", "NOT NULL, UNIQUE"],
            [3, "customer_id", "INT", "Khách hàng thành viên tích điểm (NULL nếu vãng lai)", "FK → customers(user_id), NULL"],
            [4, "sold_by", "INT", "Nhân viên trực quầy thực hiện bán hàng", "FK → staffs(user_id), NULL"],
            [5, "cinema_id", "INT", "Cụm rạp thực hiện giao dịch", "FK → cinemas(id), NULL"],
            [6, "total_price", "DECIMAL(18,0)", "Tổng tiền thanh toán của đơn hàng", "NOT NULL, DEFAULT 0"],
            [7, "payment_method", "VARCHAR(30)", "Phương thức thanh toán (CASH, CARD, TRANSFER, VNPAY)", "NULL"],
            [8, "status", "VARCHAR(30)", "Trạng thái đơn hàng (COMPLETED, CANCELLED, VOIDED)", "NOT NULL, DEFAULT 'COMPLETED'"],
            [9, "created_at", "DATETIME", "Thời điểm tạo đơn bán hàng", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 33,
        "code": "concession_sale_items",
        "title_vn": "Bảng Chi tiết mặt hàng đơn bán quầy",
        "rows": [
            [1, "id", "INT", "Khóa chính chi tiết mặt hàng đơn bán quầy", "PK, IDENTITY(1,1)"],
            [2, "sale_id", "INT", "Đơn bán ẩm thực tại quầy", "FK → concession_sales(id), NOT NULL"],
            [3, "fnb_item_id", "INT", "Mặt hàng ẩm thực được bán", "FK → fnb_items(id), NOT NULL"],
            [4, "item_name_snapshot", "NVARCHAR(100)", "Tên mặt hàng chốt tại thời điểm bán", "NULL"],
            [5, "quantity", "INT", "Số lượng sản phẩm bán", "NOT NULL, DEFAULT 1"],
            [6, "price_snapshot", "DECIMAL(18,0)", "Đơn giá chốt tại thời điểm bán", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 34,
        "code": "concession_sale_item_options",
        "title_vn": "Bảng Chi tiết tùy chọn đơn bán quầy",
        "rows": [
            [1, "id", "INT", "Khóa chính chi tiết tùy chọn món bán quầy", "PK, IDENTITY(1,1)"],
            [2, "sale_item_id", "INT", "Dòng mặt hàng bán tại quầy tương ứng", "FK → concession_sale_items(id), NOT NULL"],
            [3, "option_group_id", "INT", "Nhóm tùy chọn của khe", "FK → fnb_option_groups(id), NOT NULL"],
            [4, "option_item_id", "INT", "Tùy chọn chi tiết được chọn", "FK → fnb_option_items(id), NOT NULL"],
            [5, "slot_label_snapshot", "NVARCHAR(100)", "Nhãn khe lựa chọn chốt lúc bán", "NULL"],
            [6, "option_name_snapshot", "NVARCHAR(100)", "Tên tùy chọn chốt lúc bán", "NULL"],
            [7, "surcharge_snapshot", "DECIMAL(18,0)", "Phụ thu chốt tại thời điểm bán", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 35,
        "code": "promotions",
        "title_vn": "Bảng Chương trình khuyến mãi",
        "rows": [
            [1, "id", "INT", "Khóa chính chương trình khuyến mãi", "PK, IDENTITY(1,1)"],
            [2, "code", "VARCHAR(50)", "Mã chương trình khuyến mãi", "NOT NULL, UNIQUE"],
            [3, "name", "NVARCHAR(255)", "Tên chương trình khuyến mãi", "NULL"],
            [4, "description", "NVARCHAR(MAX)", "Mô tả chi tiết thể lệ khuyến mãi", "NULL"],
            [5, "discount_type", "VARCHAR(30)", "Loại hình giảm giá (PERCENT, FIXED, GIFT_*)", "NOT NULL"],
            [6, "discount_value", "DECIMAL(18,0)", "Giá trị giảm (% hoặc số tiền)", "NOT NULL, DEFAULT 0"],
            [7, "start_date", "DATETIME", "Thời điểm bắt đầu áp dụng", "NULL"],
            [8, "end_date", "DATETIME", "Thời điểm kết thúc áp dụng", "NULL"],
            [9, "is_stackable", "BIT", "Cho phép cộng dồn với khuyến mãi khác", "NOT NULL, DEFAULT 0"],
            [10, "points_required", "INT", "Số điểm thưởng cần để đổi mã khuyến mãi", "NOT NULL, DEFAULT 0"],
            [11, "allow_point_redemption", "BIT", "Cho phép khách dùng điểm để đổi", "NOT NULL, DEFAULT 0"],
            [12, "min_order_value", "DECIMAL(18,0)", "Giá trị đơn hàng tối thiểu để áp dụng", "NOT NULL, DEFAULT 0"],
            [13, "applicable_movie_id", "INT", "Phim cụ thể được áp dụng (soft reference)", "NULL"],
            [14, "customer_eligibility", "VARCHAR(30)", "Đối tượng khách áp dụng (ALL, NEW_CUSTOMER)", "NOT NULL, DEFAULT 'ALL'"],
            [15, "usage_limit", "INT", "Giới hạn tổng số lượt sử dụng toàn hệ thống", "NOT NULL, DEFAULT 0"],
            [16, "used_count", "INT", "Tổng số lượt đã sử dụng", "NOT NULL, DEFAULT 0"],
            [17, "max_ticket_quantity", "INT", "Số vé tối đa được giảm trong một đơn", "NOT NULL, DEFAULT 0"],
            [18, "max_discount_amount", "DECIMAL(18,0)", "Số tiền giảm tối đa (áp dụng cho giảm %)", "NOT NULL, DEFAULT 0"],
            [19, "campaign_sent_at", "DATETIME", "Thời điểm gửi email chiến dịch gần nhất", "NULL"],
            [20, "campaign_sent_count", "INT", "Số lượng email đã gửi trong chiến dịch", "NOT NULL, DEFAULT 0"]
        ]
    },
    {
        "num": 36,
        "code": "vouchers",
        "title_vn": "Bảng Phiếu giảm giá",
        "rows": [
            [1, "id", "INT", "Khóa chính phiếu giảm giá", "PK, IDENTITY(1,1)"],
            [2, "customer_id", "INT", "Khách hàng sở hữu phiếu giảm giá", "FK → customers(user_id), NOT NULL"],
            [3, "promotion_id", "INT", "Chương trình khuyến mãi phát hành voucher", "FK → promotions(id), NOT NULL"],
            [4, "valid_until", "DATETIME", "Hạn chót sử dụng voucher", "NULL"],
            [5, "is_used", "BIT", "Trạng thái đã sử dụng", "NOT NULL, DEFAULT 0"],
            [6, "used_at", "DATETIME", "Thời điểm sử dụng voucher", "NULL"]
        ]
    },
    {
        "num": 37,
        "code": "promo_email_log",
        "title_vn": "Bảng Nhật ký gửi email khuyến mãi",
        "rows": [
            [1, "id", "INT", "Khóa chính nhật ký gửi email", "PK, IDENTITY(1,1)"],
            [2, "promotion_id", "INT", "Mã chương trình khuyến mãi (soft reference)", "NOT NULL"],
            [3, "customer_id", "INT", "Mã khách hàng nhận email (soft reference)", "NOT NULL"],
            [4, "sent_at", "DATETIME", "Thời điểm gửi email", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 38,
        "code": "promo_articles",
        "title_vn": "Bảng Bài viết khuyến mãi",
        "rows": [
            [1, "id", "INT", "Khóa chính bài viết khuyến mãi", "PK, IDENTITY(1,1)"],
            [2, "title", "NVARCHAR(255)", "Tiêu đề bài viết khuyến mãi", "NOT NULL"],
            [3, "description", "NVARCHAR(500)", "Mô tả tóm tắt bài viết", "NULL"],
            [4, "image_url", "VARCHAR(500)", "Đường dẫn ảnh bìa bài viết", "NULL"],
            [5, "content", "NVARCHAR(MAX)", "Nội dung chi tiết bài viết (HTML/Markdown)", "NULL"],
            [6, "start_date", "DATE", "Ngày bắt đầu hiển thị bài viết", "NULL"],
            [7, "end_date", "DATE", "Ngày kết thúc hiển thị bài viết", "NULL"],
            [8, "is_active", "BIT", "Trạng thái kích hoạt hiển thị", "NOT NULL, DEFAULT 1"],
            [9, "display_order", "INT", "Thứ tự ưu tiên hiển thị", "NOT NULL, DEFAULT 0"],
            [10, "created_at", "DATETIME", "Thời điểm tạo bài viết", "NOT NULL, DEFAULT GETDATE()"]
        ]
    },
    {
        "num": 39,
        "code": "banners",
        "title_vn": "Bảng Banner quảng cáo",
        "rows": [
            [1, "id", "INT", "Khóa chính banner", "PK, IDENTITY(1,1)"],
            [2, "title", "NVARCHAR(255)", "Tiêu đề banner", "NULL"],
            [3, "image_url", "VARCHAR(500)", "Đường dẫn hình ảnh banner", "NULL"],
            [4, "mode", "VARCHAR(20)", "Chế độ banner (IMAGE: ảnh tĩnh, MOVIE: dựng từ phim)", "NOT NULL, DEFAULT 'IMAGE'"],
            [5, "movie_id", "INT", "Mã phim liên kết khi chọn mode MOVIE (soft reference)", "NULL"],
            [6, "placement", "VARCHAR(50)", "Vị trí đặt banner trên giao diện", "NULL"],
            [7, "start_date", "DATETIME", "Thời điểm bắt đầu hiển thị", "NULL"],
            [8, "end_date", "DATETIME", "Thời điểm kết thúc hiển thị", "NULL"],
            [9, "is_active", "BIT", "Trạng thái kích hoạt hiển thị", "NOT NULL, DEFAULT 1"],
            [10, "display_order", "INT", "Thứ tự ưu tiên hiển thị banner", "NOT NULL, DEFAULT 0"],
            [11, "link", "VARCHAR(500)", "Đường dẫn liên kết khi người dùng nhấp vào", "NULL"]
        ]
    },
    {
        "num": 40,
        "code": "pricing_rules",
        "title_vn": "Bảng Quy tắc tính giá vé",
        "rows": [
            [1, "id", "INT", "Khóa chính quy tắc giá vé", "PK, IDENTITY(1,1)"],
            [2, "name", "NVARCHAR(100)", "Tên quy tắc giá vé", "NOT NULL"],
            [3, "rule_type", "VARCHAR(30)", "Loại quy tắc (BASE_PRICE, SURCHARGE)", "NOT NULL"],
            [4, "day_type", "VARCHAR(30)", "Loại ngày áp dụng (WEEKDAY, WEDNESDAY, WEEKEND, HOLIDAY, ALL)", "NULL"],
            [5, "room_type", "VARCHAR(30)", "Loại phòng chiếu (STANDARD, SUPERPLEX, CINE_COMFORT, ALL)", "NULL"],
            [6, "time_slot", "VARCHAR(30)", "Khung giờ chiếu (EARLY, BEFORE_17H, AFTER_17H, ALL)", "NULL"],
            [7, "audience_type", "VARCHAR(30)", "Đối tượng áp dụng (ADULT, STUDENT, CHILD, SENIOR, ALL)", "NULL"],
            [8, "value", "DECIMAL(18,0)", "Mức giá áp dụng", "NOT NULL, DEFAULT 0"],
            [9, "priority", "INT", "Độ ưu tiên của quy tắc (số lớn hơn ưu tiên trước)", "NOT NULL, DEFAULT 0"],
            [10, "active", "BIT", "Trạng thái hiệu lực của quy tắc", "NOT NULL, DEFAULT 1"],
            [11, "start_date", "DATETIME", "Thời điểm bắt đầu có hiệu lực", "NULL"],
            [12, "end_date", "DATETIME", "Thời điểm kết thúc hiệu lực", "NULL"]
        ]
    },
    {
        "num": 41,
        "code": "holidays",
        "title_vn": "Bảng Ngày lễ",
        "rows": [
            [1, "id", "INT", "Khóa chính ngày lễ", "PK, IDENTITY(1,1)"],
            [2, "holiday_date", "DATE", "Ngày lễ dương lịch", "NOT NULL, UNIQUE"],
            [3, "name", "NVARCHAR(100)", "Tên ngày lễ", "NOT NULL"]
        ]
    },
    {
        "num": 42,
        "code": "faqs",
        "title_vn": "Bảng Câu hỏi thường gặp",
        "rows": [
            [1, "id", "INT", "Khóa chính câu hỏi FAQ", "PK, IDENTITY(1,1)"],
            [2, "category", "NVARCHAR(50)", "Nhóm chủ đề câu hỏi", "NOT NULL"],
            [3, "question", "NVARCHAR(500)", "Nội dung câu hỏi thường gặp", "NOT NULL"],
            [4, "answer", "NVARCHAR(MAX)", "Nội dung câu trả lời", "NULL"],
            [5, "display_order", "INT", "Thứ tự hiển thị câu hỏi", "NOT NULL, DEFAULT 0"],
            [6, "is_active", "BIT", "Trạng thái kích hoạt hiển thị", "NOT NULL, DEFAULT 1"]
        ]
    },
    {
        "num": 43,
        "code": "system_settings",
        "title_vn": "Bảng Cấu hình hệ thống",
        "rows": [
            [1, "setting_key", "VARCHAR(100)", "Khóa cấu hình hệ thống", "PK"],
            [2, "setting_value", "NVARCHAR(MAX)", "Giá trị cấu hình hệ thống", "NULL"]
        ]
    }
]

def build_docx():
    doc = docx.Document()
    
    # Page setup (A4, margins 2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(30, 30, 30)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("TỪ ĐIỂN DỮ LIỆU VẬT LÝ (PHYSICAL DATABASE SCHEMA)")
    run_title.bold = True
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(31, 78, 120)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(14)
    run_sub = subtitle_p.add_run("HỆ THỐNG QUẢN LÝ CỤM RẠP CHIẾU PHIM DEVCINE (CSDL SQL SERVER - VER9)")
    run_sub.bold = True
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(89, 89, 89)
    
    # Overview Note
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_after = Pt(16)
    run_note = note_p.add_run("Tài liệu đặc tả chi tiết 43 bảng dữ liệu mức vật lý cho hệ quản trị cơ sở dữ liệu Microsoft SQL Server, bao gồm tên trường, kiểu dữ liệu vật lý, mô tả ngữ nghĩa và các ràng buộc toàn vẹn (PK, FK, DEFAULT, UNIQUE, IDENTITY).")
    run_note.italic = True
    run_note.font.name = 'Times New Roman'
    run_note.font.size = Pt(10.5)
    run_note.font.color.rgb = RGBColor(80, 80, 80)
    
    # Loop over 43 tables
    col_widths = [Inches(0.55), Inches(1.55), Inches(1.20), Inches(2.15), Inches(1.50)]
    headers = ["STT", "Tên trường", "Kiểu dữ liệu", "Mô tả", "Ràng buộc"]
    
    for item in tables_data:
        num = item["num"]
        code = item["code"]
        title_vn = item["title_vn"]
        
        # Heading 1 line: B.x_Bảng: [tên_bảng_tiếng_Anh]
        h1_p = doc.add_paragraph()
        h1_p.paragraph_format.space_before = Pt(12)
        h1_p.paragraph_format.space_after = Pt(2)
        h1_p.paragraph_format.keep_with_next = True
        r1 = h1_p.add_run(f"B.{num}_Bảng: {code}")
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12.5)
        r1.font.color.rgb = RGBColor(31, 78, 120)
        
        # Line 2: Bảng B.x: Bảng [Tên tiếng Việt]
        h2_p = doc.add_paragraph()
        h2_p.paragraph_format.space_before = Pt(0)
        h2_p.paragraph_format.space_after = Pt(6)
        h2_p.paragraph_format.keep_with_next = True
        r2 = h2_p.add_run(f"Bảng B.{num}: {title_vn}")
        r2.bold = True
        r2.italic = True
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11.5)
        r2.font.color.rgb = RGBColor(46, 64, 83)
        
        # Table
        table = doc.add_table(rows=len(item["rows"]) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table, color="B0C4DE", sz="4")
        
        # Header row
        hdr_cells = table.rows[0].cells
        make_row_header(table.rows[0])
        make_row_cant_split(table.rows[0])
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            cell.width = col_widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, "1F4E78")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        # Data rows
        for r_idx, row_data in enumerate(item["rows"]):
            row_cells = table.rows[r_idx + 1].cells
            make_row_cant_split(table.rows[r_idx + 1])
            bg_color = "F2F7FA" if r_idx % 2 == 1 else "FFFFFF"
            
            for c_idx, val in enumerate(row_data):
                cell = row_cells[c_idx]
                cell.width = col_widths[c_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                # Alignment
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                r = p.add_run(str(val))
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                
                if c_idx == 1:
                    r.bold = True
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(15, 32, 67)
                elif c_idx == 2:
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(128, 0, 0)
                elif c_idx == 4:
                    if "PK" in str(val) or "FK" in str(val):
                        r.bold = True
                        r.font.color.rgb = RGBColor(20, 80, 150)
                    else:
                        r.font.color.rgb = RGBColor(60, 60, 60)
                else:
                    r.font.color.rgb = RGBColor(30, 30, 30)
                    
        # Add space after table
        space_p = doc.add_paragraph()
        space_p.paragraph_format.space_before = Pt(4)
        space_p.paragraph_format.space_after = Pt(8)
        
    out_dir = r"e:\DATN\DevCine\docs"
    out_path = os.path.join(out_dir, "Tu_Dien_Du_Lieu_Vat_Ly_DevCine.docx")
    doc.save(out_path)
    print(f"Generated successfully: {out_path}")

if __name__ == "__main__":
    build_docx()
