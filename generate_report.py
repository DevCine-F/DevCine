# -*- coding: utf-8 -*-
"""Sinh báo cáo thống kê & phân tích tiến độ dự án DevCine ra file Word."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ----- Bảng màu -----
PRIMARY = RGBColor(0xC8, 0x9B, 0x10)   # vàng DevCine
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
RED = RGBColor(0xB0, 0x2A, 0x2A)
ORANGE = RGBColor(0xC0, 0x6A, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Font mặc định
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align='left', white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = WHITE
    elif color:
        run.font.color.rgb = color


def add_heading(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY
        p.space_before = Pt(14)
        # gạch chân bằng border dưới
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'C89B10')
        pbdr.append(bottom)
        pPr.append(pbdr)
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        return p


def add_para(text, size=11, bold=False, color=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, color=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        if color:
            r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p


def style_header_row(row, fill='1A1A1A'):
    for cell in row.cells:
        shade_cell(cell, fill)


# ============================================================
# TRANG BÌA
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DEVCINE')
r.bold = True
r.font.size = Pt(48)
r.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HỆ THỐNG QUẢN LÝ & ĐẶT VÉ XEM PHIM')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = DARK

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO THỐNG KÊ CHỨC NĂNG\nVÀ PHÂN TÍCH TIẾN ĐỘ HOÀN THÀNH DỰ ÁN')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Công nghệ: Spring Boot 4 (Java 21) · Vue.js 3 · PostgreSQL · JWT · VNPAY')
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Ngày lập báo cáo: ' + datetime.date.today().strftime('%d/%m/%Y'))
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# 1. GIỚI THIỆU & KIẾN TRÚC
# ============================================================
add_heading('1. GIỚI THIỆU TỔNG QUAN', 1)
add_para('DevCine là hệ thống website quản lý và đặt vé xem phim trực tuyến (fullstack), '
         'phục vụ đồng thời hai nhóm người dùng: Khách hàng (đặt vé online) và Quản trị viên/Nhân viên '
         '(vận hành rạp). Hệ thống được xây dựng theo kiến trúc tách rời Frontend – Backend, giao tiếp qua REST API.')

add_heading('Kiến trúc & công nghệ sử dụng', 2)
tech = doc.add_table(rows=1, cols=2)
tech.style = 'Light Grid Accent 1'
tech.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tech.rows[0]
set_cell_text(hdr.cells[0], 'Thành phần', bold=True, white=True, size=11)
set_cell_text(hdr.cells[1], 'Công nghệ', bold=True, white=True, size=11)
style_header_row(hdr)
tech_rows = [
    ('Backend', 'Java 21, Spring Boot 4, Spring Security, Spring Data JPA, Hibernate 7'),
    ('Frontend', 'Vue.js 3.5, Pinia, Vue Router, Vite, TailwindCSS 4, Axios'),
    ('Cơ sở dữ liệu', 'PostgreSQL (hosted trên Supabase)'),
    ('Xác thực', 'JWT (JJWT 0.12.6) + BCrypt password hashing'),
    ('Thanh toán', 'Cổng VNPAY + Ví điện tử nội bộ (Wallet)'),
    ('Lưu trữ ảnh', 'Cloudinary (upload poster, banner)'),
]
for a, b in tech_rows:
    row = tech.add_row()
    set_cell_text(row.cells[0], a, bold=True, size=10)
    set_cell_text(row.cells[1], b, size=10)

doc.add_paragraph()

# ============================================================
# 2. THỐNG KÊ TỔNG QUAN
# ============================================================
add_heading('2. THỐNG KÊ TỔNG QUAN HỆ THỐNG', 1)
add_para('Sau khi rà soát toàn bộ mã nguồn, hệ thống hiện có quy mô như sau:')

stat = doc.add_table(rows=1, cols=2)
stat.style = 'Light Grid Accent 1'
hdr = stat.rows[0]
set_cell_text(hdr.cells[0], 'Hạng mục', bold=True, white=True)
set_cell_text(hdr.cells[1], 'Số lượng', bold=True, white=True, align='center')
style_header_row(hdr)
stat_rows = [
    ('Controller (REST API) ở Backend', '26'),
    ('Service (tầng nghiệp vụ)', '11'),
    ('Entity (bảng dữ liệu)', '34'),
    ('Repository (truy vấn DL)', '30'),
    ('Màn hình (View) phía Khách hàng', '19'),
    ('Màn hình (View) phía Quản trị', '20'),
    ('Tổng số nhóm chức năng nghiệp vụ', '~38'),
]
for a, b in stat_rows:
    row = stat.add_row()
    set_cell_text(row.cells[0], a, size=10)
    set_cell_text(row.cells[1], b, size=10, align='center', bold=True)

doc.add_paragraph()

# ============================================================
# 3. CHỨC NĂNG ĐÃ HOÀN THIỆN
# ============================================================
add_heading('3. CÁC CHỨC NĂNG ĐÃ HOÀN THIỆN', 1)
add_para('Là các chức năng đã kết nối hoàn chỉnh Frontend ↔ Backend ↔ Database, '
         'xử lý nghiệp vụ thực tế (không còn dữ liệu giả lập).', italic=True, color=GREY)

# --- 3.1 Nhóm khách hàng ---
add_heading('3.1. Phía Khách hàng', 2)
done_customer = [
    ('Đăng ký / Đăng nhập (JWT)',
     'Đăng ký tài khoản (tự tạo User + Customer + Wallet), đăng nhập sinh JWT token, mã hoá mật khẩu BCrypt.',
     'Có thể bổ sung: xác thực email, đăng nhập Google (nút đã có giao diện nhưng chưa nối OAuth).'),
    ('Trang chủ & Danh sách phim',
     'Hiển thị phim đang chiếu / sắp chiếu lấy động từ API.',
     'Đã đủ dùng.'),
    ('Lịch chiếu & Hệ thống rạp',
     'Lọc suất chiếu theo phim, theo thành phố/cụm rạp; lấy dữ liệu thật từ backend.',
     'Đã đủ dùng.'),
    ('Chi tiết phim',
     'Trang chi tiết phim lấy thông tin từ API theo ID.',
     'Cần bổ sung phần đánh giá/bình luận (xem mục chưa hoàn thiện).'),
    ('Đặt vé (chọn ghế + F&B + voucher)',
     'Sơ đồ ghế thời gian thực, giữ ghế (hold) có timeout 10 phút, chọn combo bắp nước, áp mã giảm giá, '
     'tính giá cuối phía server.',
     'Đã đủ dùng; có thể thêm WebSocket để khoá ghế realtime giữa nhiều người.'),
    ('Thanh toán (VNPAY + Ví điện tử)',
     'Tích hợp cổng VNPAY (tạo URL + xác thực chữ ký HMAC khi callback), thanh toán bằng ví nội bộ, '
     'cộng điểm tích luỹ & tự động nâng hạng thành viên sau khi thanh toán.',
     'Đã đủ dùng; có thể thêm Momo/ZaloPay.'),
    ('Hồ sơ cá nhân',
     'Xem & chỉnh sửa thông tin (họ tên, SĐT, ngày sinh) lưu xuống DB; hiển thị thẻ thành viên & điểm tích luỹ.',
     'Đã đủ dùng; có thể thêm đổi ảnh đại diện.'),
    ('Lịch sử đặt vé',
     'Liệt kê các đơn đã đặt kèm phim, ghế, mã vé, QR; lọc theo sắp chiếu/đã xem.',
     'Đã đủ dùng.'),
    ('Đổi mật khẩu',
     'Xác minh mật khẩu cũ, kiểm tra trùng khớp, mã hoá mật khẩu mới.',
     'Đã đủ dùng.'),
    ('Ví điện tử (Wallet)',
     'Tự tạo ví khi thiếu, nạp tiền, ghi lịch sử giao dịch, trừ tiền khi thanh toán.',
     'Phần nạp tiền thật qua cổng cần kiểm thử thêm.'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Light List Accent 1'
hdr = t.rows[0]
set_cell_text(hdr.cells[0], 'Chức năng', bold=True, white=True)
set_cell_text(hdr.cells[1], 'Đã làm được gì', bold=True, white=True)
set_cell_text(hdr.cells[2], 'Cần bổ sung thêm', bold=True, white=True)
style_header_row(hdr, '1E7E34')
for name, did, more in done_customer:
    row = t.add_row()
    set_cell_text(row.cells[0], name, bold=True, size=9.5)
    set_cell_text(row.cells[1], did, size=9.5)
    set_cell_text(row.cells[2], more, size=9.5, color=GREY)

doc.add_paragraph()

# --- 3.2 Nhóm quản trị ---
add_heading('3.2. Phía Quản trị / Nhân viên', 2)
done_admin = [
    ('Đăng nhập Admin (JWT + phân vai)',
     'Đăng nhập riêng cho quản trị, chỉ cho phép vai trò ADMIN/STAFF truy cập.',
     'Cần siết phân quyền chi tiết bằng @PreAuthorize (xem mục chưa hoàn thiện).'),
    ('Dashboard tổng quan',
     'Thống kê doanh thu, số vé, số liệu vận hành lấy từ DashboardService.',
     'Có thể thêm biểu đồ theo khoảng thời gian tuỳ chọn.'),
    ('Quản lý phim',
     'CRUD phim đầy đủ, upload poster qua Cloudinary.',
     'Đã đủ dùng.'),
    ('Quản lý danh mục phim',
     'Quản lý thể loại, gắn thể loại cho phim.',
     'Đã đủ dùng.'),
    ('Quản lý Banner quảng cáo',
     'CRUD banner trang chủ, bật/tắt hiển thị, sắp thứ tự (vừa hoàn thiện kết nối API).',
     'Đã đủ dùng.'),
    ('Điều phối lịch chiếu (Master Scheduling)',
     'Tạo/sắp xếp suất chiếu theo phòng và khung giờ, có kiểm tra trùng lịch.',
     'Đã đủ dùng.'),
    ('Quản lý cụm rạp & phòng chiếu',
     'CRUD rạp, phòng, cấu hình; có tab phân tích.',
     'Đã đủ dùng.'),
    ('Sơ đồ ghế (Seat Map Editor)',
     'Thiết kế ma trận ghế theo phòng, gán loại ghế (thường/VIP/sweetbox).',
     'Đã đủ dùng.'),
    ('Quản lý nhân sự & Ca trực',
     'Danh sách nhân viên, xếp ca, duyệt/từ chối ca trực theo ngày.',
     'Đã đủ dùng; có thể thêm bàn giao ca (ShiftHandover entity đã có nhưng chưa dùng).'),
    ('Bán vé tại quầy (POS)',
     'Bán vé trực tiếp: chọn suất, ghế, combo, tra thẻ thành viên, thanh toán.',
     'Đã đủ dùng.'),
    ('Kiểm soát vé / Check-in (QR)',
     'Quét QR bằng camera hoặc nhập tay, xác thực vé, chống check-in trùng & vé quá hạn, có âm báo.',
     'Đã đủ dùng.'),
    ('Khuyến mãi & Phát voucher',
     'CRUD chương trình khuyến mãi, phát voucher cho khách theo promotion.',
     'Đã đủ dùng.'),
    ('Cấu hình giá (Pricing Rule)',
     'CRUD quy tắc tính giá theo khung giờ/loại ghế/ngày.',
     'Đã đủ dùng.'),
    ('Quản lý kho F&B (Inventory)',
     'Danh sách hàng hoá, nhập/điều chỉnh tồn kho, ghi log kho.',
     'Có thể nối với BomRecipe để tự trừ nguyên liệu khi bán combo.'),
    ('Chăm sóc khách hàng (Support Ticket)',
     'Tiếp nhận yêu cầu hỗ trợ, cập nhật trạng thái Chờ → Đang xử lý → Đóng (vừa hoàn thiện).',
     'Có thể thêm chức năng phản hồi/nhắn lại cho khách.'),
    ('Cài đặt hệ thống',
     'Lưu cấu hình: tên site, email, giá vé cơ bản, tỉ lệ điểm thưởng, chế độ bảo trì.',
     'Đã đủ dùng.'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Light List Accent 1'
hdr = t.rows[0]
set_cell_text(hdr.cells[0], 'Chức năng', bold=True, white=True)
set_cell_text(hdr.cells[1], 'Đã làm được gì', bold=True, white=True)
set_cell_text(hdr.cells[2], 'Cần bổ sung thêm', bold=True, white=True)
style_header_row(hdr, '1E7E34')
for name, did, more in done_admin:
    row = t.add_row()
    set_cell_text(row.cells[0], name, bold=True, size=9.5)
    set_cell_text(row.cells[1], did, size=9.5)
    set_cell_text(row.cells[2], more, size=9.5, color=GREY)

doc.add_paragraph()

# ============================================================
# 4. CHỨC NĂNG CHƯA HOÀN THIỆN
# ============================================================
add_heading('4. CÁC CHỨC NĂNG CHƯA HOÀN THIỆN', 1)
add_para('Là các chức năng đã có giao diện nhưng còn dùng dữ liệu giả lập, hoặc thiếu hẳn tầng xử lý '
         'phía backend, chưa hoạt động thực tế.', italic=True, color=GREY)

incomplete = [
    ('Đánh giá & bình luận phim',
     'Thiếu toàn bộ tầng xử lý: đã có Entity "Review" trong CSDL nhưng KHÔNG có Controller/Service/Repository, '
     'giao diện chi tiết phim chỉ hiển thị tĩnh.',
     'Khách đã xem phim có thể chấm sao (1–5) và viết nhận xét; hiển thị điểm trung bình và danh sách bình luận '
     'trên trang chi tiết phim; admin kiểm duyệt/ẩn bình luận vi phạm.'),
    ('Nhật ký hệ thống (Audit Log)',
     'Màn hình "Nhật ký" ĐÃ đọc được dữ liệu qua API, NHƯNG không có đoạn mã nào GHI log: '
     'entity AuditLog hiện là "mồ côi" nên bảng log luôn rỗng.',
     'Mỗi thao tác quan trọng (tạo/sửa/xoá phim, suất chiếu, hoàn vé, đăng nhập admin…) cần tự động ghi 1 bản ghi '
     'AuditLog (ai làm, hành động gì, thời điểm, IP). Nên dùng AOP hoặc gọi trong từng service.'),
    ('Phân quyền chi tiết (Permissions)',
     'Màn hình phân quyền dùng dữ liệu cứng (mock), không lưu xuống DB và backend chưa enforce quyền theo từng '
     'chức năng (mới chỉ chặn theo vai trò chung ADMIN/STAFF).',
     'Cho phép gán quyền (xem/thêm/sửa/xoá/xuất) theo từng module cho từng vai trò, lưu DB; '
     'backend kiểm tra quyền bằng @PreAuthorize/đối tượng quyền trước mỗi API.'),
    ('Voucher của tôi (phía khách)',
     'Trang "Ưu đãi của tôi" đang hiển thị danh sách voucher GIẢ LẬP, dù backend đã có VoucherController '
     '(API lấy voucher theo khách đã sẵn sàng).',
     'Chỉ cần nối API /api/vouchers/customer/{id} để hiển thị voucher thật của khách (đang dùng / đã dùng / hết hạn).'),
    ('Thông báo (Notification)',
     'Trang thông báo hoàn toàn tĩnh; chưa có Entity/Controller cho thông báo.',
     'Sinh thông báo cho khách (đặt vé thành công, nhắc giờ chiếu, ưu đãi mới); '
     'có badge số chưa đọc, đánh dấu đã đọc. Cần thêm bảng Notification + API.'),
    ('Tìm kiếm phim',
     'Trang tìm kiếm chỉ là giao diện tĩnh (luôn hiện trạng thái "không tìm thấy"), chưa gọi API tìm kiếm.',
     'Ô tìm kiếm gọi API lọc phim theo tên/thể loại/đạo diễn, hiển thị kết quả thật và gợi ý liên quan.'),
    ('Trang Khuyến mãi (phía khách)',
     'Trang "Khuyến mãi & Ưu đãi" là nội dung tĩnh, các nút "Learn More" trỏ tới "#".',
     'Lấy danh sách chương trình khuyến mãi đang chạy từ backend (đã có Promotion); '
     'mỗi ưu đãi mở trang chi tiết, có thể "lưu mã" vào ví voucher của khách.'),
    ('Lịch chiếu (màn hình AdminSchedule cũ)',
     'Màn hình lịch chiếu này dùng dữ liệu mock và tính năng "AI Auto-Scheduling" chỉ là thanh tiến trình giả lập. '
     '(Phần điều phối thật đã chuyển sang Master Scheduling).',
     'Hoặc gỡ bỏ màn hình trùng lặp này, hoặc nối dữ liệu thật và hiện thực thuật toán xếp lịch tự động '
     '(tối ưu doanh thu/khung giờ vàng).'),
    ('Thanh toán (trang PaymentView độc lập)',
     'Trang /payment là form thẻ tín dụng tĩnh, không xử lý. Luồng thanh toán thật đã nằm trong màn Đặt vé.',
     'Nên gỡ bỏ hoặc chuyển hướng về luồng đặt vé chuẩn để tránh gây nhầm lẫn.'),
    ('Bàn giao ca & Định mức nguyên liệu',
     'Hai Entity ShiftHandover và BomRecipe đã khai báo nhưng chưa có nghiệp vụ sử dụng.',
     'ShiftHandover: ghi nhận tiền quỹ/ghi chú khi đổi ca. BomRecipe: định mức nguyên liệu cho mỗi combo để '
     'tự trừ tồn kho khi bán hàng.'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Light List Accent 1'
hdr = t.rows[0]
set_cell_text(hdr.cells[0], 'Chức năng', bold=True, white=True)
set_cell_text(hdr.cells[1], 'Đang thiếu gì', bold=True, white=True)
set_cell_text(hdr.cells[2], 'Nên hoạt động như thế nào', bold=True, white=True)
style_header_row(hdr, 'B02A2A')
for name, miss, should in incomplete:
    row = t.add_row()
    set_cell_text(row.cells[0], name, bold=True, size=9.5)
    set_cell_text(row.cells[1], miss, size=9.5, color=RED)
    set_cell_text(row.cells[2], should, size=9.5, color=GREY)

doc.add_paragraph()

# ============================================================
# 5. SO SÁNH & TIẾN ĐỘ
# ============================================================
add_heading('5. SO SÁNH & TIẾN ĐỘ HOÀN THÀNH', 1)

add_heading('5.1. Bảng tiến độ theo nhóm chức năng', 2)
prog = doc.add_table(rows=1, cols=4)
prog.style = 'Light Grid Accent 1'
hdr = prog.rows[0]
set_cell_text(hdr.cells[0], 'Nhóm chức năng', bold=True, white=True)
set_cell_text(hdr.cells[1], 'Đã xong', bold=True, white=True, align='center')
set_cell_text(hdr.cells[2], 'Chưa xong', bold=True, white=True, align='center')
set_cell_text(hdr.cells[3], 'Tỉ lệ', bold=True, white=True, align='center')
style_header_row(hdr)
prog_rows = [
    ('Khách hàng', 12, 6, '67%'),
    ('Quản trị / Nhân viên', 16, 3, '84%'),
    ('Hạ tầng (Auth/Security/Thanh toán)', 4, 1, '80%'),
]
for name, done, undone, rate in prog_rows:
    row = prog.add_row()
    set_cell_text(row.cells[0], name, bold=True, size=10)
    set_cell_text(row.cells[1], str(done), size=10, align='center', color=GREEN, bold=True)
    set_cell_text(row.cells[2], str(undone), size=10, align='center', color=RED, bold=True)
    set_cell_text(row.cells[3], rate, size=10, align='center', bold=True)
# Tổng
row = prog.add_row()
set_cell_text(row.cells[0], 'TỔNG CỘNG', bold=True, size=10.5)
set_cell_text(row.cells[1], '32', size=10.5, align='center', color=GREEN, bold=True)
set_cell_text(row.cells[2], '10', size=10.5, align='center', color=RED, bold=True)
set_cell_text(row.cells[3], '~78%', size=10.5, align='center', bold=True, color=PRIMARY)
for c in row.cells:
    shade_cell(c, 'F2E9C8')

doc.add_paragraph()

add_heading('5.2. Đánh giá tiến độ tổng thể', 2)
add_para('Mức độ hoàn thành ước tính: ', bold=True, space_after=2)
p = doc.add_paragraph()
r = p.add_run('≈ 78% ')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = PRIMARY
r2 = p.add_run('toàn dự án')
r2.font.size = Pt(12)
r2.font.color.rgb = GREY

add_para('Toàn bộ luồng nghiệp vụ CỐT LÕI đã hoàn chỉnh và chạy thật: '
         'duyệt phim → chọn suất → đặt vé/chọn ghế → thanh toán (VNPAY/Ví) → sinh vé QR → check-in tại rạp, '
         'cùng các nghiệp vụ vận hành quan trọng (POS, kho, nhân sự, khuyến mãi, định giá, dashboard).',
         space_after=6)
add_para('Phần còn lại (~22%) chủ yếu là các chức năng PHỤ TRỢ và bổ sung trải nghiệm: '
         'đánh giá phim, thông báo, tìm kiếm, phân quyền chi tiết, ghi nhật ký, và một số trang tĩnh '
         'cần nối dữ liệu (voucher, khuyến mãi).', space_after=6)

# Thanh tiến độ trực quan (bằng bảng 1 hàng 100 ô tô màu)
bar = doc.add_table(rows=1, cols=10)
bar.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, cell in enumerate(bar.rows[0].cells):
    if i < 8:  # ~78% -> ~8/10
        shade_cell(cell, 'C89B10')
    else:
        shade_cell(cell, 'E0E0E0')
    cell.text = ''
    cell.width = Inches(0.6)

doc.add_paragraph()

# ============================================================
# 6. KHUYẾN NGHỊ LỘ TRÌNH
# ============================================================
add_heading('6. KHUYẾN NGHỊ LỘ TRÌNH HOÀN THIỆN', 1)
add_para('Đề xuất thứ tự ưu tiên để đưa dự án về mức hoàn chỉnh:', space_after=6)

add_heading('Ưu tiên cao (ảnh hưởng trực tiếp nghiệp vụ & bảo mật)', 2)
add_bullet('Hiện thực ghi Audit Log cho mọi thao tác quản trị quan trọng (để màn Nhật ký có dữ liệu thật).',
           color=RED, bold_prefix='[1] ')
add_bullet('Siết phân quyền chi tiết ở backend bằng @PreAuthorize + lưu cấu hình quyền xuống DB.',
           color=RED, bold_prefix='[2] ')
add_bullet('Nối trang "Voucher của tôi" với VoucherController (API đã sẵn — chi phí thấp, giá trị cao).',
           color=RED, bold_prefix='[3] ')

add_heading('Ưu tiên trung bình (hoàn thiện trải nghiệm khách hàng)', 2)
add_bullet('Xây dựng module Đánh giá & bình luận phim (Controller + Service + UI sao/nhận xét).',
           color=ORANGE, bold_prefix='[4] ')
add_bullet('Hiện thực Tìm kiếm phim thật và nối trang Khuyến mãi với dữ liệu Promotion.',
           color=ORANGE, bold_prefix='[5] ')
add_bullet('Xây dựng hệ thống Thông báo (bảng Notification + API + badge chưa đọc).',
           color=ORANGE, bold_prefix='[6] ')

add_heading('Ưu tiên thấp (dọn dẹp & mở rộng)', 2)
add_bullet('Gỡ/hợp nhất màn hình lịch chiếu cũ (AdminSchedule) và trang PaymentView tĩnh để tránh trùng lặp.',
           color=GREY, bold_prefix='[7] ')
add_bullet('Hiện thực Bàn giao ca (ShiftHandover) và Định mức nguyên liệu (BomRecipe) để tự trừ kho.',
           color=GREY, bold_prefix='[8] ')
add_bullet('Bổ sung đăng nhập Google (OAuth) và mở rộng cổng thanh toán (Momo/ZaloPay).',
           color=GREY, bold_prefix='[9] ')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— HẾT BÁO CÁO —')
r.italic = True
r.font.color.rgb = GREY

# Lưu file
out = r'E:\DATN\DevCine\Bao_cao_thong_ke_DevCine.docx'
doc.save(out)
print('DONE:', out)
